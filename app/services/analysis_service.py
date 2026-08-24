
"""
Candidate Analysis Orchestration
================================

This service is the application boundary between the Flask project workspace
and the existing intelligence pipeline.

Main flow:

    Resume DOCX + JD DOCX
            |
            v
    DocumentProcessingService
            |
            v
    DocumentKnowledgeProfile + JDRequirementProfile
            |
            v
    Phase 3 / 4 Matching
            |
            v
    Phase 5 ATS Analysis
            |
            v
    Candidate-facing Result
            |
            v
    Recommendations + Repository/Ontology Gap Report

Important:
- No resume rewriting is performed here.
- Cover letters are intentionally outside this pipeline.
- The service never fabricates a missing skill as a candidate capability.
- Missing repository / ontology terms are reported for administrator review.
"""

from __future__ import annotations

import importlib
import json
import logging
import re

from dataclasses import asdict, fields, is_dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from app.intelligence.utilities.knowledge.documents.document_input import (
    DocumentInput,
)
from app.intelligence.utilities.knowledge.documents.document_types import (
    DocumentType,
)
from app.intelligence.utilities.knowledge.documents.document_profile_builder import (
    DocumentProfileBuilder,
)
from app.intelligence.utilities.knowledge.documents.document_processing.document_processing_service import (
    DocumentProcessingService,
)
from app.intelligence.utilities.knowledge.jd_requirements.requirement_classifier import (
    JDRequirementClassifier,
)

from app.parser.resume_parser import ResumeParser
from app.parser.resume_builder import ResumeBuilder
from app.intelligence.utilities.knowledge.education.education_equivalence import (
    EducationEquivalence,
)


logger = logging.getLogger("GetHired")


class AnalysisPipelineError(RuntimeError):
    """Raised when the candidate-analysis pipeline cannot complete."""


class CandidateAnalysisService:
    """Run and persist the candidate-vs-JD analysis."""

    # Keys that may be produced by ontology/repository components when
    # something was not recognized or is not currently present in the
    # repository.
    REPOSITORY_GAP_KEYS = {
        "unknown_terms",
        "unknown_entities",
        "unmatched_terms",
        "unrecognized_terms",
        "repository_gaps",
        "ontology_gaps",
        "missing_entities",
        "missing_ontology",
        "missing_ontology_terms",
        "unresolved_entities",
        "unresolved_terms",
        "unmapped_entities",
        "unmapped_terms",
    }

    # Experience-role evidence is deliberately kept outside ontology matching.
    # These title families answer requirements such as:
    #   "previous experience in a supervisory or managerial position"
    # using the candidate's structured employment history.
    EXPERIENCE_ROLE_FAMILIES = {
        "managerial_supervisory": (
            "manager",
            "managing director",
            "general manager",
            "store manager",
            "operations manager",
            "production manager",
            "plant manager",
            "quality manager",
            "qa manager",
            "food safety manager",
            "managerial",
            "director",
            "supervisor",
            "superintendent",
            "team lead",
            "team leader",
            "shift lead",
            "shift supervisor",
            "head of",
            "department head",
            "chief",
            "president",
            "vice president",
            "vp ",
            "owner",
            "founder",
        ),
        "managerial": (
            "manager",
            "managing director",
            "general manager",
            "director",
            "managerial",
            "chief",
            "president",
            "vice president",
            "vp ",
            "owner",
            "founder",
        ),
        "supervisory": (
            "supervisor",
            "manager",
            "store manager",
            "team lead",
            "team leader",
            "shift lead",
            "shift supervisor",
            "foreman",
            "head of",
            "department head",
        ),
        "leadership": (
            "manager",
            "managing director",
            "director",
            "supervisor",
            "team lead",
            "team leader",
            "head of",
            "chief",
            "president",
            "vice president",
            "vp ",
            "owner",
            "founder",
        ),
        "executive": (
            "managing director",
            "director",
            "general manager",
            "chief",
            "president",
            "vice president",
            "vp ",
            "owner",
            "founder",
            "ceo",
            "coo",
            "cfo",
            "cto",
        ),
    }

    def __init__(self, app_config: dict[str, Any]):
        self.upload_folder = Path(app_config["UPLOAD_FOLDER"])
        self.output_folder = Path(app_config["OUTPUT_FOLDER"])

        self.document_service = DocumentProcessingService()
        self.profile_builder = DocumentProfileBuilder()
        self.jd_classifier = JDRequirementClassifier()
        self.resume_parser = ResumeParser()
        self.resume_builder = ResumeBuilder()

    # ==================================================================
    # PUBLIC API
    # ==================================================================

    def run(self, project_id: int) -> dict[str, Any]:
        """
        Run the complete candidate analysis.

        The analysis is intentionally triggered explicitly by the user from
        the project workspace rather than automatically after file upload.
        """

        project_uploads = (
            self.upload_folder / f"project_{project_id}"
        )

        project_outputs = (
            self.output_folder / f"project_{project_id}"
        )

        project_outputs.mkdir(
            parents=True,
            exist_ok=True,
        )

        resume_path = (
            project_uploads / "resume_original.docx"
        )

        jd_path = (
            project_uploads / "job_description_original.docx"
        )

        # --------------------------------------------------------------
        # Validate input files
        # --------------------------------------------------------------

        if not resume_path.exists():
            raise AnalysisPipelineError(
                "Resume file is missing."
            )

        if not jd_path.exists():
            raise AnalysisPipelineError(
                "Job Description file is missing."
            )

        logger.info(
            "Analysis started | project=%s",
            project_id,
        )

        # --------------------------------------------------------------
        # Read DOCX files
        # --------------------------------------------------------------

        resume_text = self._read_docx(
            resume_path
        )

        jd_text = self._read_docx(
            jd_path
        )

        logger.info(
            "Documents loaded | project=%s | resume_chars=%s | jd_chars=%s",
            project_id,
            len(resume_text),
            len(jd_text),
        )

        structured_resume = self._build_structured_resume_evidence(
            resume_path
        )

        # --------------------------------------------------------------
        # Resume document pipeline
        # --------------------------------------------------------------

        resume_response = self.document_service.process(
            DocumentInput(
                text=resume_text,
                document_type=DocumentType.RESUME,
            )
        )

        if not resume_response.success:
            raise AnalysisPipelineError(
                "Resume pipeline failed: "
                f"{resume_response.error}"
            )

        logger.info(
            "Resume processing completed | project=%s",
            project_id,
        )

        # --------------------------------------------------------------
        # JD document pipeline
        # --------------------------------------------------------------

        jd_response = self.document_service.process(
            DocumentInput(
                text=jd_text,
                document_type=DocumentType.JD,
            )
        )

        if not jd_response.success:
            raise AnalysisPipelineError(
                "Job Description pipeline failed: "
                f"{jd_response.error}"
            )

        logger.info(
            "JD processing completed | project=%s",
            project_id,
        )

        # --------------------------------------------------------------
        # Build profiles
        # --------------------------------------------------------------

        resume_profile = self.profile_builder.build(
            resume_response
        )

        jd_profile = self.profile_builder.build(
            jd_response
        )

        logger.info(
            "Document profiles built | project=%s",
            project_id,
        )

        # --------------------------------------------------------------
        # Classify JD requirements
        # --------------------------------------------------------------

        jd_requirements = self.jd_classifier.process(
            jd_profile
        )

        logger.info(
            "JD requirements classified | project=%s",
            project_id,
        )

        # --------------------------------------------------------------
        # Phase 3 / Phase 4 matching
        # --------------------------------------------------------------

        matching = self._build_matching(
            resume_profile=resume_profile,
            jd_profile=jd_profile,
            jd_requirements=jd_requirements,
            resume_text=resume_text,
            structured_resume=structured_resume,
        )

        logger.info(
            "Matching completed | project=%s | overall=%s",
            project_id,
            matching["overall_match"],
        )

        # --------------------------------------------------------------
        # Phase 5 ATS analysis
        # --------------------------------------------------------------

        ats = self._run_latest_ats_if_available(
            resume_text=resume_text,
            matching=matching,
            jd_requirements=jd_requirements,
            structured_resume=structured_resume,
        )

        logger.info(
            "ATS analysis completed | project=%s | score=%s",
            project_id,
            ats["score"],
        )

        # --------------------------------------------------------------
        # Build candidate-facing result
        # --------------------------------------------------------------

        result = self._build_candidate_result(
            project_id=project_id,
            resume_text=resume_text,
            jd_text=jd_text,
            resume_response=resume_response,
            jd_response=jd_response,
            resume_profile=resume_profile,
            jd_profile=jd_profile,
            jd_requirements=jd_requirements,
            matching=matching,
            ats=ats,
            structured_resume=structured_resume,
        )

        # --------------------------------------------------------------
        # Repository / ontology gaps
        # --------------------------------------------------------------

        repository_gaps = self._collect_repository_gaps(
            resume_response.result,
            jd_response.result,
            matching.get(
                "repository_gaps",
                [],
            ),
        )

        result["repository_gaps"] = repository_gaps

        result["completed_at"] = (
            datetime.now(timezone.utc).isoformat()
        )

        # --------------------------------------------------------------
        # Persist project result
        # --------------------------------------------------------------

        self._write_json(
            project_outputs / "analysis_result.json",
            result,
        )

        self._write_json(
            project_outputs / "repository_gaps.json",
            {
                "project_id": project_id,
                "generated_at": result["completed_at"],
                "gaps": repository_gaps,
            },
        )

        self._write_repository_gap_text(
            project_outputs / "repository_gaps.txt",
            project_id,
            repository_gaps,
        )

        # --------------------------------------------------------------
        # Persist cumulative administrator report
        # --------------------------------------------------------------

        self._update_admin_repository_report(
            project_id=project_id,
            gaps=repository_gaps,
        )

        logger.info(
            "Analysis completed | project=%s | overall=%s | ats=%s | repository_gaps=%s",
            project_id,
            result["scores"]["overall_match"],
            result["scores"]["ats_compatibility"],
            len(repository_gaps),
        )

        return result

    def load(
        self,
        project_id: int,
    ) -> dict[str, Any] | None:
        """Load a previously completed project analysis."""

        path = (
            self.output_folder
            / f"project_{project_id}"
            / "analysis_result.json"
        )

        if not path.exists():
            return None

        try:
            return json.loads(
                path.read_text(
                    encoding="utf-8"
                )
            )

        except (
            OSError,
            json.JSONDecodeError,
        ):
            logger.exception(
                "Could not load analysis result | project=%s",
                project_id,
            )

            return None

    # ==================================================================
    # DOCUMENT INPUT
    # ==================================================================

    @staticmethod
    def _read_docx(
        path: Path,
    ) -> str:
        """
        Read paragraphs and tables from a DOCX document.
        """

        try:
            document = Document(
                str(path)
            )

        except Exception as exc:
            raise AnalysisPipelineError(
                f"Could not read DOCX '{path.name}': {exc}"
            ) from exc

        parts: list[str] = []

        # --------------------------------------------------------------
        # Paragraphs
        # --------------------------------------------------------------

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # --------------------------------------------------------------
        # Tables
        # --------------------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                cells = [
                    cell
                    for cell in cells
                    if cell
                ]

                if cells:
                    parts.append(
                        " | ".join(cells)
                    )

        text = "\n".join(parts).strip()

        if not text:
            raise AnalysisPipelineError(
                f"DOCX '{path.name}' contains no readable text."
            )

        return text

    # ==================================================================
    # MATCHING
    # ==================================================================

    def _build_matching(
        self,
        *,
        resume_profile: Any,
        jd_profile: Any,
        jd_requirements: Any,
        resume_text: str,
        structured_resume: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """
        Build Phase 3 / Phase 4 matching objects when available.

        If the newer typed matcher is not installed, the local bridge is
        used so the front end can still receive useful results.
        """

        resume_entities = list(
            getattr(
                getattr(
                    resume_profile,
                    "entities",
                    None,
                ),
                "entities",
                [],
            )
            or []
        )

        requirements = list(
            getattr(
                jd_requirements,
                "requirements",
                (),
            )
            or ()
        )

        records: list[dict[str, Any]] = []

        for requirement in requirements:

            records.append(
                self._match_requirement(
                    requirement,
                    resume_entities,
                    resume_text,
                    structured_resume or {},
                )
            )

        typed_profile = (
            self._try_build_typed_match_profile(
                records=records,
                resume_profile=resume_profile,
                jd_requirements=jd_requirements,
            )
        )

        strong = [
            item
            for item in records
            if item["status"] == "matched"
        ]

        partial = [
            item
            for item in records
            if item["status"] == "partial"
        ]

        missing = [
            item
            for item in records
            if item["status"] == "unmatched"
        ]

        # --------------------------------------------------------------
        # Weighted overall match
        # --------------------------------------------------------------

        weighted_total = 0.0
        weight_total = 0.0

        for item in records:

            weight = self._priority_weight(
                item["priority"]
            )

            weighted_total += (
                item["score"] * weight
            )

            weight_total += weight

        overall = (
            weighted_total / weight_total
            if weight_total
            else 0.0
        )

        # --------------------------------------------------------------
        # Evidence strength
        # --------------------------------------------------------------

        evidence_values = [
            item["evidence_confidence"]
            for item in records
            if item["status"] != "unmatched"
        ]

        evidence_strength = (
            sum(evidence_values)
            / len(evidence_values)
            if evidence_values
            else 0.0
        )

        return {
            "requirements": records,

            "strong_matches": strong,

            "partial_matches": partial,

            "missing_matches": missing,

            "overall_match": round(
                overall * 100,
                1,
            ),

            "evidence_strength": round(
                evidence_strength * 100,
                1,
            ),

            "typed_profile": typed_profile,

            "repository_gaps": [],
        }

    def _try_build_typed_match_profile(
        self,
        *,
        records: list[dict[str, Any]],
        resume_profile: Any,
        jd_requirements: Any,
    ) -> Any:
        """
        Use the newer typed Phase 3 / Phase 4 implementation when it exists.
        """

        try:

            match_models = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.match_models"
            )

            enrichment_models = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.enrichment_models"
            )

            enricher_module = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.match_enricher"
            )

            gap_module = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.gap_analyzer"
            )

            profile_builder_module = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.knowledge_match_profile_builder"
            )

            # Keep import alive for compatibility with repositories where
            # enrichment_models performs registration at import time.
            _ = enrichment_models

            requirement_match_cls = (
                match_models.RequirementMatch
            )

            match_result_cls = (
                match_models.KnowledgeMatchResult
            )

            match_status = (
                match_models.MatchStatus
            )

            match_basis = (
                match_models.MatchBasis
            )

            typed_matches = []

            for record in records:

                status_name = record["status"]

                status = {
                    "matched": getattr(
                        match_status,
                        "MATCHED",
                    ),
                    "partial": getattr(
                        match_status,
                        "PARTIAL",
                    ),
                    "unmatched": getattr(
                        match_status,
                        "UNMATCHED",
                    ),
                }[status_name]

                basis = (
                    getattr(
                        match_basis,
                        "CANONICAL",
                    )
                    if status_name in {
                        "matched",
                        "partial",
                    }
                    else getattr(
                        match_basis,
                        "NONE",
                    )
                )

                typed_matches.append(
                    requirement_match_cls(
                        requirement_id=record[
                            "requirement_id"
                        ],
                        requirement_subject=record[
                            "subject"
                        ],
                        requirement_type=record[
                            "requirement_type"
                        ],
                        priority=record[
                            "priority"
                        ],
                        status=status,
                        score=record[
                            "score"
                        ],
                        basis=basis,
                        candidate_entity_ids=tuple(
                            record.get(
                                "candidate_entity_ids",
                                [],
                            )
                        ),
                        candidate_evidence=tuple(
                            record.get(
                                "evidence",
                                [],
                            )
                        ),
                        evidence_count=len(
                            record.get(
                                "evidence",
                                [],
                            )
                        ),
                    )
                )

            match_result = (
                match_result_cls.from_matches(
                    typed_matches
                )
            )

            enriched = (
                enricher_module
                .KnowledgeMatchEnricher()
                .process(
                    match_result=match_result,
                    resume_profile=resume_profile,
                    jd_requirement_profile=jd_requirements,
                )
            )

            gap = (
                gap_module
                .KnowledgeGapAnalyzer()
                .process(
                    enriched
                )
            )

            profile = (
                profile_builder_module
                .KnowledgeMatchProfileBuilder()
                .process(
                    match_result=match_result,
                    enriched_match_result=enriched,
                    gap_analysis_result=gap,
                )
            )

            return profile

        except (
            ImportError,
            AttributeError,
            TypeError,
            ValueError,
        ) as exc:

            logger.info(
                "Typed Phase 3/4 matcher unavailable; "
                "using application bridge: %s",
                exc,
            )

            return None

    def _match_requirement(
        self,
        requirement: Any,
        resume_entities: list[Any],
        resume_text: str,
        structured_resume: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """
        Match one JD requirement against resume entities.
        """

        subject = str(
            getattr(
                requirement,
                "subject",
                "",
            )
            or ""
        ).strip()

        requirement_type = self._enum_value(
            getattr(
                requirement,
                "requirement_type",
                "unknown",
            )
        )

        priority = self._enum_value(
            getattr(
                requirement,
                "priority",
                "contextual",
            )
        )

        subject_norm = self._normalize(
            subject
        )

        subject_tokens = set(
            subject_norm.split()
        )

        # --------------------------------------------------------------
        # Education equivalency
        # --------------------------------------------------------------
        # Education is intentionally handled outside ontology matching. A
        # higher academic level satisfies a lower-level requirement when the
        # JD does not impose a conflicting field constraint. For example, an
        # M.Sc. satisfies a Bachelor's degree requirement.
        # --------------------------------------------------------------
        if requirement_type == "education":
            education_records = (structured_resume or {}).get(
                "education",
                [],
            )
            education_match = EducationEquivalence.match_requirement(
                requirement,
                education_records,
            )

            evidence = []
            if education_match.candidate_degree:
                candidate_description = education_match.candidate_degree
                if education_match.candidate_major:
                    candidate_description += (
                        f" ({education_match.candidate_major})"
                    )
                evidence.append(
                    "Structured education evidence: "
                    + candidate_description
                )

            if education_match.reason:
                evidence.append(education_match.reason)

            return {
                "requirement_id": getattr(requirement, "requirement_id", ""),
                "subject": subject,
                "requirement_type": requirement_type,
                "priority": priority,
                "minimum_years": getattr(requirement, "minimum_years", None),
                "status": "matched" if education_match.matched else (
                    "partial" if education_match.score >= 0.45 else "unmatched"
                ),
                "score": round(education_match.score, 4),
                "basis": "education_equivalence",
                "candidate_entity_ids": [],
                "evidence": evidence,
                "evidence_confidence": round(
                    0.96 if education_match.matched else max(0.55, education_match.score),
                    4,
                ),
                "education_match": {
                    "required_level": education_match.required_level,
                    "candidate_level": education_match.candidate_level,
                    "candidate_degree": education_match.candidate_degree,
                    "candidate_major": education_match.candidate_major,
                    "equivalent": education_match.equivalent,
                    "field_required": education_match.field_required,
                    "field_matched": education_match.field_matched,
                    "reason": education_match.reason,
                },
            }

        candidates = []

        for entity in resume_entities:

            canonical = self._entity_text(
                entity,
                "canonical",
                "label",
                "normalized",
                "name",
                "original",
            )

            if not canonical:
                continue

            entity_norm = self._normalize(
                canonical
            )

            if not entity_norm:
                continue

            entity_tokens = set(
                entity_norm.split()
            )

            exact = (
                subject_norm
                == entity_norm
            )

            substring = (
                subject_norm in entity_norm
                or entity_norm in subject_norm
            )

            overlap = (
                len(
                    subject_tokens
                    & entity_tokens
                )
                / max(
                    len(
                        subject_tokens
                        | entity_tokens
                    ),
                    1,
                )
            )

            score = 0.0
            basis = "none"

            if exact:

                score = 1.0
                basis = "canonical"

            elif substring:

                score = 0.82
                basis = "related"

            elif overlap >= 0.5:

                score = 0.65
                basis = "semantic_overlap"

            if score > 0:

                candidates.append(
                    (
                        score,
                        basis,
                        entity,
                    )
                )

        candidates.sort(
            key=lambda item: item[0],
            reverse=True,
        )

        best = (
            candidates[0]
            if candidates
            else None
        )

        if best is None:

            status = "unmatched"
            score = 0.0
            basis = "none"
            candidate = None

        elif best[0] >= 0.8:

            status = "matched"
            score = best[0]
            basis = best[1]
            candidate = best[2]

        else:

            status = "partial"
            score = best[0]
            basis = best[1]
            candidate = best[2]

        evidence: list[str] = []
        candidate_entity_ids: list[str] = []
        evidence_confidence = 0.0

        if candidate is not None:

            candidate_id = self._entity_text(
                candidate,
                "entity_id",
                "id",
                "node_id",
            )

            if candidate_id:
                candidate_entity_ids.append(
                    candidate_id
                )

            source_text = self._entity_text(
                candidate,
                "source_text",
                "description",
                "achievement",
                "evidence",
            )

            if source_text:
                evidence.append(
                    source_text
                )

            try:

                evidence_confidence = float(
                    getattr(
                        candidate,
                        "confidence",
                        0.0,
                    )
                    if not isinstance(
                        candidate,
                        dict,
                    )
                    else candidate.get(
                        "confidence",
                        0.0,
                    )
                )

            except (
                TypeError,
                ValueError,
            ):

                evidence_confidence = 0.0

        # --------------------------------------------------------------
        # Non-ontology experience-role evidence
        # --------------------------------------------------------------
        # A requirement such as "previous experience in a supervisory or
        # managerial position" is not a literal keyword requirement.  It must
        # be evaluated against employment titles.  Ontology overlap alone can
        # produce false positives (for example Manager -> managerial) without
        # proving that the candidate actually held that role.
        if requirement_type == "experience":
            role_match = self._experience_role_match(
                subject=subject,
                structured_resume=structured_resume or {},
                resume_text=resume_text,
            )
            if role_match is not None:
                role_score = float(role_match.get("score", 0.0))
                if role_score >= 0.82:
                    status = "matched"
                    score = max(score, role_score)
                elif role_score >= 0.55 and status != "matched":
                    status = "partial"
                    score = max(score, role_score)
                evidence.append(role_match["evidence"])
                evidence_confidence = max(
                    evidence_confidence,
                    float(role_match.get("confidence", 0.0)),
                )

                minimum_years = getattr(requirement, "minimum_years", None)
                qualifying_years = role_match.get("qualifying_years")
                if minimum_years is not None and qualifying_years is not None:
                    minimum = float(minimum_years)
                    if qualifying_years >= minimum:
                        status = "matched"
                        score = max(score, 1.0)
                        evidence.append(
                            f"Qualifying supervisory/managerial roles total approximately "
                            f"{qualifying_years:g} years; JD asks for {minimum:g}+ years."
                        )
                    elif qualifying_years > 0 and status != "matched":
                        status = "partial"
                        score = max(score, min(qualifying_years / minimum, 0.79))
                        evidence.append(
                            f"Qualifying supervisory/managerial roles total approximately "
                            f"{qualifying_years:g} years; JD asks for {minimum:g}+ years."
                        )

        # --------------------------------------------------------------
        # Non-ontology structured candidate evidence
        # --------------------------------------------------------------
        structured_candidates = self._structured_resume_candidates(
            requirement_type=requirement_type,
            subject=subject,
            structured_resume=structured_resume or {},
        )
        if structured_candidates:
            best_structured = structured_candidates[0]
            structured_score = best_structured["score"]
            if structured_score >= 0.82:
                status = "matched"
                score = max(score, structured_score)
            elif structured_score >= 0.55 and status != "matched":
                status = "partial"
                score = max(score, structured_score)
            evidence.append(best_structured["evidence"])
            evidence_confidence = max(
                evidence_confidence,
                float(best_structured.get("confidence", 0.0)),
            )

        # --------------------------------------------------------------
        # Experience requirements
        # --------------------------------------------------------------

        if requirement_type == "experience":

            minimum_years = getattr(
                requirement,
                "minimum_years",
                None,
            )

            resume_years = (
                self._resume_years(
                    resume_text
                )
            )

            if (
                minimum_years is not None
                and resume_years is not None
            ):

                if (
                    resume_years
                    >= float(minimum_years)
                ):

                    status = "matched"

                    score = max(
                        score,
                        1.0,
                    )

                    evidence.append(
                        "Resume contains approximately "
                        f"{resume_years:g} years "
                        "of stated experience."
                    )

                elif resume_years > 0:

                    status = "partial"

                    score = max(
                        score,
                        min(
                            resume_years
                            / float(
                                minimum_years
                            ),
                            0.79,
                        ),
                    )

                    evidence.append(
                        "Resume states approximately "
                        f"{resume_years:g} years; "
                        "JD asks for "
                        f"{float(minimum_years):g}+ years."
                    )

        return {
            "requirement_id": str(
                getattr(
                    requirement,
                    "requirement_id",
                    subject,
                )
            ),
            "subject": subject,
            "requirement_type": requirement_type,
            "priority": priority,
            "mandatory": bool(
                getattr(
                    requirement,
                    "mandatory",
                    False,
                )
            ),
            "preferred": bool(
                getattr(
                    requirement,
                    "preferred",
                    False,
                )
            ),
            "minimum_years": getattr(
                requirement,
                "minimum_years",
                None,
            ),
            "status": status,
            "score": round(
                score,
                4,
            ),
            "basis": basis,
            "candidate_entity_ids": candidate_entity_ids,
            "evidence": evidence,
            "evidence_confidence": round(
                evidence_confidence,
                4,
            ),
        }

    # ==================================================================
    # ATS
    # ==================================================================

    def _run_latest_ats_if_available(
        self,
        *,
        resume_text: str,
        matching: dict[str, Any],
        jd_requirements: Any = None,
        structured_resume: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """
        Run the latest Phase 5 ATS analyzer when available.
        Otherwise use the safe local fallback.
        """

        profile = matching.get(
            "typed_profile"
        )

        if profile is not None:

            try:

                policy_module = importlib.import_module(
                    "app.intelligence.utilities.knowledge.ats.ats_analysis_policy"
                )

                analyzer_module = importlib.import_module(
                    "app.intelligence.utilities.knowledge.ats.ats_resume_analyzer"
                )

                request_module = importlib.import_module(
                    "app.intelligence.utilities.knowledge.ats.ats_analysis_request"
                )

                policy = (
                    policy_module
                    .ATSAnalysisPolicy()
                )

                analyzer = (
                    analyzer_module
                    .ATSResumeAnalyzer(
                        policy=policy
                    )
                )

                request = (
                    request_module
                    .ATSResumeAnalysisRequest(
                        resume_text=resume_text,
                        knowledge_match_profile=profile,
                        resume_profile=getattr(profile, "resume_profile", None),
                        jd_requirement_profile=jd_requirements,
                        metadata={
                            "structured_resume": structured_resume or {},
                            "jd_requirement_profile": jd_requirements,
                        },
                    )
                )

                result = analyzer.process(
                    request
                )

                score = float(
                    getattr(
                        getattr(
                            result,
                            "ats_score",
                            None,
                        ),
                        "score",
                        getattr(
                            getattr(
                                result,
                                "score",
                                None,
                            ),
                            "score",
                            0.0,
                        ),
                    )
                )

                component_fields = {
                    "keyword_analysis": "keyword_coverage_score",
                    "section_analysis": "section_completeness_score",
                    "formatting_analysis": "formatting_score",
                    "readability_analysis": "readability_score",
                    "terminology_analysis": "terminology_score",
                    "quantification_analysis": "quantification_score",
                    "parseability_analysis": "parseability_score",
                }

                ui_components: dict[str, dict[str, Any]] = {}
                for component_name, score_field in component_fields.items():
                    ui_components[component_name] = self._ats_component_for_ui(
                        getattr(result, component_name, None),
                        score_field,
                    )

                return {
                    "available": True,
                    "score": round(score * 100, 1),
                    "confidence": round(
                        float(getattr(result, "confidence", 0.0)) * 100,
                        1,
                    ),
                    **ui_components,
                }

            except (
                ImportError,
                AttributeError,
                TypeError,
                ValueError,
            ) as exc:

                logger.info(
                    "Latest Phase 5 ATS unavailable: %s",
                    exc,
                )

        return self._fallback_ats(
            resume_text,
            matching,
        )

    def _fallback_ats(
        self,
        resume_text: str,
        matching: dict[str, Any],
    ) -> dict[str, Any]:
        """
        Safe ATS fallback used only when the full ATS analyzer is unavailable.
        """

        requirements = matching[
            "requirements"
        ]

        required_subjects = [
            item["subject"]
            for item in requirements
            if item["priority"]
            in {
                "required",
                "high",
            }
        ]

        if not required_subjects:

            required_subjects = [
                item["subject"]
                for item in requirements
            ]

        lower = resume_text.casefold()

        matched_keywords = [
            subject
            for subject in required_subjects
            if subject.casefold()
            in lower
        ]

        missing_keywords = [
            subject
            for subject in required_subjects
            if subject.casefold()
            not in lower
        ]

        keyword_score = (
            len(matched_keywords)
            / len(required_subjects)
            if required_subjects
            else 1.0
        )

        required_sections = (
            "summary",
            "experience",
            "skills",
            "education",
        )

        detected_sections = [
            section
            for section in required_sections
            if section in lower
        ]

        section_score = (
            len(detected_sections)
            / len(required_sections)
        )

        quantified = len(
            re.findall(
                r"\b\d+(?:\.\d+)?\s*%",
                resume_text,
            )
        )

        quantification_score = min(
            quantified / 3.0,
            1.0,
        )

        parseability_score = (
            1.0
            if resume_text.strip()
            else 0.0
        )

        score = (
            keyword_score * 0.55
            + section_score * 0.20
            + quantification_score * 0.15
            + parseability_score * 0.10
        )

        return {
            "available": False,
            "score": round(
                score * 100,
                1,
            ),
            "confidence": 60.0,
            "keyword_analysis": {
                "required_keywords": required_subjects,
                "matched_keywords": matched_keywords,
                "missing_keywords": missing_keywords,
                "keyword_coverage_score": round(
                    keyword_score * 100,
                    1,
                ),
            },
            "section_analysis": {
                "detected_sections": detected_sections,
                "missing_sections": [
                    item
                    for item in required_sections
                    if item not in detected_sections
                ],
                "section_completeness_score": round(
                    section_score * 100,
                    1,
                ),
            },
            "quantification_analysis": {
                "quantified_achievement_count": quantified,
                "quantification_score": round(
                    quantification_score * 100,
                    1,
                ),
            },
            "parseability_analysis": {
                "parseable": bool(
                    resume_text.strip()
                ),
                "parseability_score": round(
                    parseability_score * 100,
                    1,
                ),
            },
        }

    # ==================================================================
    # RESULT
    # ==================================================================

    def _build_candidate_result(
        self,
        **kwargs: Any,
    ) -> dict[str, Any]:

        resume_profile = kwargs[
            "resume_profile"
        ]

        jd_profile = kwargs[
            "jd_profile"
        ]

        jd_requirements = kwargs[
            "jd_requirements"
        ]

        matching = kwargs[
            "matching"
        ]

        ats = kwargs[
            "ats"
        ]

        resume_response = kwargs[
            "resume_response"
        ]

        jd_response = kwargs[
            "jd_response"
        ]

        structured_resume = kwargs.get(
            "structured_resume",
            {},
        )

        # --------------------------------------------------------------
        # Experience gaps
        # --------------------------------------------------------------

        experience_gap = []

        for item in matching[
            "requirements"
        ]:

            if (
                item["requirement_type"]
                == "experience"
            ):

                experience_gap.append(
                    {
                        "jd_requirement": item[
                            "subject"
                        ],
                        "minimum_years": item[
                            "minimum_years"
                        ],
                        "status": item[
                            "status"
                        ],
                        "evidence": item[
                            "evidence"
                        ],
                    }
                )

        recommendations = (
            self._recommendations(
                matching=matching,
                ats=ats,
            )
        )

        return {
            "project_id": kwargs[
                "project_id"
            ],

            "scores": {
                "overall_match": matching[
                    "overall_match"
                ],
                "ats_compatibility": ats[
                    "score"
                ],
                "evidence_strength": matching[
                    "evidence_strength"
                ],
            },

            "matches": {
                "strong": matching[
                    "strong_matches"
                ],
                "partial": matching[
                    "partial_matches"
                ],
                "missing": matching[
                    "missing_matches"
                ],
            },

            "experience_gap": experience_gap,

            "recommendations": recommendations,

            "categories": self._category_summary(
                resume_profile,
                jd_profile,
            ),

            "requirements_summary": {
                "total": len(
                    getattr(
                        jd_requirements,
                        "requirements",
                        (),
                    )
                    or ()
                ),
                "required": getattr(
                    jd_requirements,
                    "required_count",
                    0,
                ),
                "preferred": getattr(
                    jd_requirements,
                    "preferred_count",
                    0,
                ),
                "contextual": getattr(
                    jd_requirements,
                    "contextual_count",
                    0,
                ),
            },

            "jd_requirement_details": [
                {
                    "requirement_id": getattr(item, "requirement_id", ""),
                    "type": self._enum_value(getattr(item, "requirement_type", "unknown")),
                    "priority": self._enum_value(getattr(item, "priority", "contextual")),
                    "subject": getattr(item, "subject", ""),
                    "minimum_years": getattr(item, "minimum_years", None),
                    "evidence": getattr(item, "evidence", ""),
                    "section": (getattr(item, "metadata", {}) or {}).get("section", ""),
                    "evidence_kind": (getattr(item, "metadata", {}) or {}).get("evidence_kind", "ontology"),
                    "confidence": getattr(item, "confidence", 0.0),
                }
                for item in (getattr(jd_requirements, "requirements", ()) or ())
            ],

            "structured_resume": structured_resume,

            "ats": ats,

            "pipeline": {
                "resume_entities": len(
                    getattr(
                        getattr(
                            resume_profile,
                            "entities",
                            None,
                        ),
                        "entities",
                        [],
                    )
                    or []
                ),

                "jd_entities": len(
                    getattr(
                        getattr(
                            jd_profile,
                            "entities",
                            None,
                        ),
                        "entities",
                        [],
                    )
                    or []
                ),

                "resume_semantic_entities": len(
                    getattr(
                        resume_response,
                        "semantic_entities",
                        [],
                    )
                    or []
                ),

                "jd_semantic_entities": len(
                    getattr(
                        jd_response,
                        "semantic_entities",
                        [],
                    )
                    or []
                ),

                "resume_business_statements": len(
                    getattr(
                        resume_response,
                        "business_statements",
                        [],
                    )
                    or []
                ),

                "jd_business_statements": len(
                    getattr(
                        jd_response,
                        "business_statements",
                        [],
                    )
                    or []
                ),

                "typed_phase4_profile": (
                    matching[
                        "typed_profile"
                    ]
                    is not None
                ),
            },
        }

    # ==================================================================
    # RECOMMENDATIONS
    # ==================================================================

    def _recommendations(
        self,
        *,
        matching: dict[str, Any],
        ats: dict[str, Any],
    ) -> list[str]:

        recommendations: list[str] = []

        # --------------------------------------------------------------
        # Strong matches
        # --------------------------------------------------------------

        for item in matching[
            "strong_matches"
        ][:3]:

            if item["evidence"]:

                recommendations.append(
                    "Surface "
                    f"{item['subject']} "
                    "evidence clearly in the resume."
                )

        # --------------------------------------------------------------
        # Partial matches
        # --------------------------------------------------------------

        for item in matching[
            "partial_matches"
        ][:4]:

            recommendations.append(
                "Strengthen evidence for "
                f"{item['subject']} "
                "if the experience is genuinely possessed."
            )

        # --------------------------------------------------------------
        # Missing matches
        # --------------------------------------------------------------

        for item in matching[
            "missing_matches"
        ][:5]:

            recommendations.append(
                "Consider "
                f"{item['subject']} "
                "only if genuinely possessed; "
                "do not add unsupported keywords."
            )

        # --------------------------------------------------------------
        # ATS sections
        # --------------------------------------------------------------

        missing_sections = (
            ats.get(
                "section_analysis",
                {},
            )
            or {}
        ).get(
            "missing_sections",
            [],
        )

        if missing_sections:

            recommendations.append(
                "Improve resume structure by "
                "addressing missing ATS-relevant sections."
            )

        # --------------------------------------------------------------
        # Quantification
        # --------------------------------------------------------------

        quantification = (
            ats.get(
                "quantification_analysis",
                {},
            )
            or {}
        ).get(
            "quantification_score",
            100,
        )

        if float(
            quantification or 0
        ) < 70:

            recommendations.append(
                "Add measurable outcomes to "
                "genuine achievements where the "
                "resume currently lacks quantified evidence."
            )

        if not recommendations:

            recommendations.append(
                "The current analysis found no "
                "high-confidence improvement recommendation "
                "beyond maintaining evidence accuracy."
            )

        # --------------------------------------------------------------
        # De-duplicate
        # --------------------------------------------------------------

        return list(
            dict.fromkeys(
                recommendations
            )
        )

    # ==================================================================
    # CATEGORIES
    # ==================================================================

    def _category_summary(
        self,
        resume_profile: Any,
        jd_profile: Any,
    ) -> dict[str, list[str]]:

        entities = list(
            getattr(
                getattr(
                    resume_profile,
                    "entities",
                    None,
                ),
                "entities",
                [],
            )
            or []
        )

        jd_entities = list(
            getattr(
                getattr(
                    jd_profile,
                    "entities",
                    None,
                ),
                "entities",
                [],
            )
            or []
        )

        groups = {
            "skills": {
                "skill"
            },
            "standards": {
                "standard"
            },
            "certifications": {
                "certification"
            },
            "methodologies": {
                "methodology"
            },
            "technologies": {
                "technology"
            },
            "actions": {
                "action"
            },
            "domains": {
                "domain"
            },
            "metrics": {
                "metric",
                "kpi",
            },
        }

        output: dict[
            str,
            list[str],
        ] = {}

        for label, types in groups.items():

            output[label] = sorted(
                self._unique(
                    self._entity_text(
                        entity,
                        "canonical",
                        "label",
                        "name",
                        "normalized",
                    )
                    for entity in entities
                    if self._enum_value(
                        self._entity_text(
                            entity,
                            "entity_type",
                            "type",
                            "category",
                        )
                    )
                    in types
                )
            )

        output[
            "jd_requirements"
        ] = sorted(
            self._unique(
                self._entity_text(
                    entity,
                    "canonical",
                    "label",
                    "name",
                    "normalized",
                )
                for entity in jd_entities
            )
        )

        return output

    # ==================================================================
    # REPOSITORY / ONTOLOGY GAP DETECTION
    # ==================================================================

    def _collect_repository_gaps(
        self,
        *sources: Any,
    ) -> list[dict[str, Any]]:
        """
        Recursively inspect pipeline output for explicit repository /
        ontology gap fields.

        The important distinction is:

        - We DO NOT guess that a word is missing from the repository.
        - We ONLY report terms when an upstream component explicitly exposes
          them through known gap/unresolved fields.

        This prevents ordinary candidate skills from incorrectly becoming
        "repository gaps".
        """

        found: list[
            dict[str, Any]
        ] = []

        for index, source in enumerate(
            sources
        ):

            source_name = (
                f"source_{index + 1}"
            )

            self._walk_repository_gaps(
                value=source,
                target=found,
                source=source_name,
                path=(),
                seen=set(),
                depth=0,
            )

        # --------------------------------------------------------------
        # De-duplicate
        # --------------------------------------------------------------

        unique: dict[
            tuple[str, str],
            dict[str, Any],
        ] = {}

        for item in found:

            term = str(
                item.get(
                    "term",
                    "",
                )
                or ""
            ).strip()

            entity_type = str(
                item.get(
                    "entity_type",
                    "unknown",
                )
                or "unknown"
            ).strip()

            if not term:
                continue

            key = (
                term.casefold(),
                entity_type.casefold(),
            )

            if key not in unique:
                unique[key] = item

            else:
                # Preserve all source locations.
                old_sources = set(
                    unique[key].get(
                        "sources",
                        [],
                    )
                )

                old_sources.update(
                    item.get(
                        "sources",
                        [],
                    )
                )

                unique[key][
                    "sources"
                ] = sorted(
                    old_sources
                )

        result = sorted(
            unique.values(),
            key=lambda item: (
                str(
                    item.get(
                        "term",
                        "",
                    )
                ).casefold(),
                str(
                    item.get(
                        "entity_type",
                        "",
                    )
                ).casefold(),
            ),
        )

        return result

    def _walk_repository_gaps(
        self,
        *,
        value: Any,
        target: list[dict[str, Any]],
        source: str,
        path: tuple[str, ...],
        seen: set[int],
        depth: int,
    ) -> None:
        """
        Recursive repository-gap walker.

        This is deliberately conservative. It only extracts values attached
        to explicit repository/ontology gap fields.
        """

        if depth > 25:
            return

        if value is None:
            return

        # --------------------------------------------------------------
        # Primitive
        # --------------------------------------------------------------

        if isinstance(
            value,
            (
                str,
                int,
                float,
                bool,
            ),
        ):
            return

        # --------------------------------------------------------------
        # Dataclass
        # --------------------------------------------------------------

        if (
            is_dataclass(value)
            and not isinstance(
                value,
                type,
            )
        ):

            object_id = id(value)

            if object_id in seen:
                return

            seen.add(
                object_id
            )

            for field in fields(
                value
            ):

                try:
                    child = getattr(
                        value,
                        field.name,
                    )
                except Exception:
                    continue

                field_name = (
                    field.name.casefold()
                )

                child_path = (
                    path
                    + (
                        field.name,
                    )
                )

                if field_name in (
                    self.REPOSITORY_GAP_KEYS
                ):

                    self._append_gap_values(
                        target=target,
                        value=child,
                        source=source,
                        path=child_path,
                    )

                else:

                    self._walk_repository_gaps(
                        value=child,
                        target=target,
                        source=source,
                        path=child_path,
                        seen=seen,
                        depth=depth + 1,
                    )

            return

        # --------------------------------------------------------------
        # Dictionary
        # --------------------------------------------------------------

        if isinstance(
            value,
            dict,
        ):

            object_id = id(value)

            if object_id in seen:
                return

            seen.add(
                object_id
            )

            for key, child in value.items():

                key_string = str(
                    key
                )

                key_name = (
                    key_string.casefold()
                )

                child_path = (
                    path
                    + (
                        key_string,
                    )
                )

                if key_name in (
                    self.REPOSITORY_GAP_KEYS
                ):

                    self._append_gap_values(
                        target=target,
                        value=child,
                        source=source,
                        path=child_path,
                    )

                else:

                    self._walk_repository_gaps(
                        value=child,
                        target=target,
                        source=source,
                        path=child_path,
                        seen=seen,
                        depth=depth + 1,
                    )

            return

        # --------------------------------------------------------------
        # Lists / tuples / sets
        # --------------------------------------------------------------

        if isinstance(
            value,
            (
                list,
                tuple,
                set,
                frozenset,
            ),
        ):

            object_id = id(value)

            if object_id in seen:
                return

            seen.add(
                object_id
            )

            for index, child in enumerate(
                value
            ):

                child_path = (
                    path
                    + (
                        str(index),
                    )
                )

                self._walk_repository_gaps(
                    value=child,
                    target=target,
                    source=source,
                    path=child_path,
                    seen=seen,
                    depth=depth + 1,
                )

            return

        # --------------------------------------------------------------
        # Generic object
        # --------------------------------------------------------------

        if hasattr(
            value,
            "__dict__",
        ):

            object_id = id(value)

            if object_id in seen:
                return

            seen.add(
                object_id
            )

            try:
                attributes = vars(
                    value
                )
            except Exception:
                return

            for key, child in attributes.items():

                key_name = str(
                    key
                ).casefold()

                child_path = (
                    path
                    + (
                        str(key),
                    )
                )

                if key_name in (
                    self.REPOSITORY_GAP_KEYS
                ):

                    self._append_gap_values(
                        target=target,
                        value=child,
                        source=source,
                        path=child_path,
                    )

                else:

                    self._walk_repository_gaps(
                        value=child,
                        target=target,
                        source=source,
                        path=child_path,
                        seen=seen,
                        depth=depth + 1,
                    )

    @staticmethod
    def _append_gap_values(
        *,
        target: list[dict[str, Any]],
        value: Any,
        source: str,
        path: tuple[str, ...],
    ) -> None:
        """
        Convert an explicit repository-gap value into normalized records.

        Supported forms include:

            "Power BI"

            {
                "term": "Power BI",
                "entity_type": "technology"
            }

            {
                "label": "ISO 45001",
                "type": "standard"
            }

            [
                "SAP",
                "Power BI"
            ]
        """

        # --------------------------------------------------------------
        # String
        # --------------------------------------------------------------

        if isinstance(
            value,
            str,
        ):

            term = value.strip()

            if term:

                target.append(
                    {
                        "term": term,
                        "entity_type": "unknown",
                        "source": source,
                        "path": ".".join(
                            path
                        ),
                        "sources": [
                            source
                        ],
                    }
                )

            return

        # --------------------------------------------------------------
        # Dictionary
        # --------------------------------------------------------------

        if isinstance(
            value,
            dict,
        ):

            term = (
                value.get("term")
                or value.get("label")
                or value.get("canonical")
                or value.get("text")
                or value.get("name")
                or value.get("subject")
            )

            entity_type = (
                value.get(
                    "entity_type"
                )
                or value.get(
                    "type"
                )
                or value.get(
                    "category"
                )
                or "unknown"
            )

            if term:

                target.append(
                    {
                        "term": str(
                            term
                        ).strip(),
                        "entity_type": str(
                            entity_type
                        ).strip(),
                        "source": source,
                        "path": ".".join(
                            path
                        ),
                        "sources": [
                            source
                        ],
                    }
                )

                return

            # Some upstream components may return nested structures.
            for key, child in value.items():

                self_source = (
                    f"{source}:{key}"
                )

                CandidateAnalysisService._append_gap_values(
                    target=target,
                    value=child,
                    source=self_source,
                    path=path + (
                        str(key),
                    ),
                )

            return

        # --------------------------------------------------------------
        # List / tuple / set
        # --------------------------------------------------------------

        if isinstance(
            value,
            (
                list,
                tuple,
                set,
                frozenset,
            ),
        ):

            for index, item in enumerate(
                value
            ):

                CandidateAnalysisService._append_gap_values(
                    target=target,
                    value=item,
                    source=source,
                    path=path + (
                        str(index),
                    ),
                )

    # ==================================================================
    # ADMIN REPOSITORY REPORT
    # ==================================================================

    def _update_admin_repository_report(
        self,
        *,
        project_id: int,
        gaps: list[dict[str, Any]],
    ) -> None:
        """
        Update the cumulative administrator repository-gap report.

        Location:

            outputs/admin/repository_gap_report.json
            outputs/admin/repository_gap_report.txt

        Existing entries are preserved and merged.
        """

        admin_folder = (
            self.output_folder
            / "admin"
        )

        admin_folder.mkdir(
            parents=True,
            exist_ok=True,
        )

        json_path = (
            admin_folder
            / "repository_gap_report.json"
        )

        text_path = (
            admin_folder
            / "repository_gap_report.txt"
        )

        # --------------------------------------------------------------
        # Load previous report
        # --------------------------------------------------------------

        existing_entries: list[
            dict[str, Any]
        ] = []

        if json_path.exists():

            try:

                existing = json.loads(
                    json_path.read_text(
                        encoding="utf-8"
                    )
                )

                existing_entries = list(
                    existing.get(
                        "gaps",
                        [],
                    )
                    or []
                )

            except (
                OSError,
                json.JSONDecodeError,
            ):

                logger.warning(
                    "Existing admin repository report "
                    "could not be read; rebuilding it."
                )

        # --------------------------------------------------------------
        # Merge current entries
        # --------------------------------------------------------------

        merged: dict[
            tuple[str, str],
            dict[str, Any],
        ] = {}

        for item in existing_entries:

            term = str(
                item.get(
                    "term",
                    "",
                )
                or ""
            ).strip()

            entity_type = str(
                item.get(
                    "entity_type",
                    "unknown",
                )
                or "unknown"
            ).strip()

            if not term:
                continue

            key = (
                term.casefold(),
                entity_type.casefold(),
            )

            merged[key] = {
                **item,
                "projects": list(
                    item.get(
                        "projects",
                        [],
                    )
                    or []
                ),
                "sources": list(
                    item.get(
                        "sources",
                        [],
                    )
                    or []
                ),
            }

        for gap in gaps:

            term = str(
                gap.get(
                    "term",
                    "",
                )
                or ""
            ).strip()

            entity_type = str(
                gap.get(
                    "entity_type",
                    "unknown",
                )
                or "unknown"
            ).strip()

            if not term:
                continue

            key = (
                term.casefold(),
                entity_type.casefold(),
            )

            if key not in merged:

                merged[key] = {
                    "term": term,
                    "entity_type": entity_type,
                    "projects": [],
                    "sources": [],
                    "first_seen": datetime.now(
                        timezone.utc
                    ).isoformat(),
                }

            entry = merged[key]

            projects = set(
                entry.get(
                    "projects",
                    [],
                )
                or []
            )

            projects.add(
                project_id
            )

            entry["projects"] = sorted(
                projects
            )

            sources = set(
                entry.get(
                    "sources",
                    [],
                )
                or []
            )

            gap_sources = gap.get(
                "sources",
                [],
            ) or []

            if not gap_sources:

                gap_source = gap.get(
                    "source"
                )

                if gap_source:
                    gap_sources = [
                        gap_source
                    ]

            sources.update(
                str(item)
                for item in gap_sources
                if str(item).strip()
            )

            entry["sources"] = sorted(
                sources
            )

            entry[
                "last_seen"
            ] = datetime.now(
                timezone.utc
            ).isoformat()

        final_gaps = sorted(
            merged.values(),
            key=lambda item: (
                str(
                    item.get(
                        "term",
                        "",
                    )
                ).casefold(),
                str(
                    item.get(
                        "entity_type",
                        "",
                    )
                ).casefold(),
            ),
        )

        # --------------------------------------------------------------
        # JSON report
        # --------------------------------------------------------------

        report = {
            "report_name": (
                "GetHired Repository / Ontology Gap Report"
            ),
            "generated_at": datetime.now(
                timezone.utc
            ).isoformat(),
            "total_unique_gaps": len(
                final_gaps
            ),
            "gaps": final_gaps,
        }

        self._write_json(
            json_path,
            report,
        )

        # --------------------------------------------------------------
        # Human-readable admin report
        # --------------------------------------------------------------

        lines = [
            "GETHIRED REPOSITORY / ONTOLOGY GAP REPORT",
            "===========================================",
            "",
            f"Generated: {report['generated_at']}",
            f"Unique gaps: {len(final_gaps)}",
            "",
        ]

        if not final_gaps:

            lines.extend(
                [
                    "No repository / ontology gaps have been reported.",
                    "",
                ]
            )

        else:

            for index, gap in enumerate(
                final_gaps,
                start=1,
            ):

                term = gap.get(
                    "term",
                    "",
                )

                entity_type = gap.get(
                    "entity_type",
                    "unknown",
                )

                projects = gap.get(
                    "projects",
                    [],
                )

                sources = gap.get(
                    "sources",
                    [],
                )

                lines.append(
                    f"{index}. {term}"
                )

                lines.append(
                    f"   Type: {entity_type}"
                )

                lines.append(
                    "   Projects: "
                    + ", ".join(
                        str(project)
                        for project in projects
                    )
                )

                lines.append(
                    "   Sources: "
                    + ", ".join(
                        str(source)
                        for source in sources
                    )
                )

                lines.append("")

        text_path.write_text(
            "\n".join(lines),
            encoding="utf-8",
        )

        logger.info(
            "Admin repository report updated | "
            "unique_gaps=%s | path=%s",
            len(final_gaps),
            json_path,
        )

    # ==================================================================
    # SERIALIZATION
    # ==================================================================

    @staticmethod
    def _ats_component_for_ui(
        value: Any,
        score_field: str,
    ) -> dict[str, Any]:
        """Serialize one typed ATS component using frontend percentage units.

        Phase 5 typed models intentionally store scores in [0, 1].  The
        candidate-analysis JSON, however, is a presentation boundary and the
        existing frontend renders these fields with a literal ``%`` suffix.
        Convert exactly once here so a normalized 1.0 becomes 100.0%, not
        1.0%.
        """
        component = CandidateAnalysisService._object_summary(value)
        if component and score_field in component:
            try:
                raw_score = float(component[score_field])
                # Typed ATS models use [0, 1]. Older/fallback payloads may
                # already use [0, 100]. Normalize exactly once.
                if 0.0 <= raw_score <= 1.0:
                    raw_score *= 100.0
                component[score_field] = round(
                    max(0.0, min(100.0, raw_score)),
                    1,
                )
            except (TypeError, ValueError):
                pass
        return component

    @staticmethod
    def _object_summary(
        value: Any,
    ) -> dict[str, Any]:

        if value is None:
            return {}

        if is_dataclass(value):

            return CandidateAnalysisService._json_safe(
                asdict(value)
            )

        if isinstance(
            value,
            dict,
        ):

            return CandidateAnalysisService._json_safe(
                value
            )

        if hasattr(
            value,
            "__dict__",
        ):

            return CandidateAnalysisService._json_safe(
                vars(value)
            )

        return {
            "value": CandidateAnalysisService._json_safe(
                value
            )
        }

    @staticmethod
    def _json_safe(
        value: Any,
    ) -> Any:

        if value is None:
            return None

        if isinstance(
            value,
            (
                str,
                int,
                float,
                bool,
            ),
        ):

            return value

        if isinstance(
            value,
            dict,
        ):

            return {
                str(key): CandidateAnalysisService._json_safe(
                    item
                )
                for key, item in value.items()
            }

        if isinstance(
            value,
            (
                list,
                tuple,
                set,
            ),
        ):

            return [
                CandidateAnalysisService._json_safe(
                    item
                )
                for item in value
            ]

        if hasattr(
            value,
            "value",
        ) and not isinstance(
            value,
            (
                str,
                bytes,
            ),
        ):

            try:

                return CandidateAnalysisService._json_safe(
                    value.value
                )

            except Exception:
                pass

        if is_dataclass(
            value
        ):

            return CandidateAnalysisService._json_safe(
                asdict(value)
            )

        if hasattr(
            value,
            "__dict__",
        ):

            return CandidateAnalysisService._json_safe(
                vars(value)
            )

        return str(
            value
        )

    @staticmethod
    def _write_json(
        path: Path,
        value: Any,
    ) -> None:

        path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        path.write_text(
            json.dumps(
                CandidateAnalysisService._json_safe(
                    value
                ),
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

    @staticmethod
    def _write_repository_gap_text(
        path: Path,
        project_id: int,
        gaps: list[
            dict[str, Any]
        ],
    ) -> None:

        lines = [
            "GETHIRED REPOSITORY / ONTOLOGY GAP REPORT",
            "===========================================",
            f"Project: {project_id}",
            (
                "Generated: "
                f"{datetime.now(timezone.utc).isoformat()}"
            ),
            "",
        ]

        if not gaps:

            lines.append(
                "No explicit repository/ontology gaps were reported by the pipeline."
            )

        else:

            for gap in gaps:

                lines.append(
                    "- "
                    f"{gap.get('term', '')} "
                    "| type="
                    f"{gap.get('entity_type', 'unknown')} "
                    "| source="
                    f"{gap.get('source', '')}"
                )

        path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        path.write_text(
            "\n".join(lines),
            encoding="utf-8",
        )

    # ==================================================================
    # HELPERS
    # ==================================================================

    @staticmethod
    def _enum_value(
        value: Any,
    ) -> str:

        return str(
            getattr(
                value,
                "value",
                value,
            )
        ).strip().casefold()

    @staticmethod
    def _normalize(
        value: str,
    ) -> str:

        value = str(
            value or ""
        ).casefold()

        value = re.sub(
            r"[^a-z0-9+#.&/-]+",
            " ",
            value,
        )

        return re.sub(
            r"\s+",
            " ",
            value,
        ).strip()

    @staticmethod
    def _entity_text(
        entity: Any,
        *names: str,
    ) -> str:
        """
        Safely extract textual values from an entity.

        Important:
        ----------
        Boolean values such as False must NEVER become textual evidence.

        Without this protection:

            evidence=False

        becomes:

            "False"

        which then gets displayed as resume evidence.

        This helper is used for canonical/entity/source/evidence fields,
        so structured values must be converted conservatively.
        """

        # --------------------------------------------------------------
        # DICTIONARY ENTITY
        # --------------------------------------------------------------

        if isinstance(entity, dict):

            for name in names:

                value = entity.get(name)

                # Boolean values are not textual evidence.
                if isinstance(value, bool):
                    continue

                if (
                    value is not None
                    and str(value).strip()
                ):

                    return str(
                        value
                    ).strip()

            return ""

        # --------------------------------------------------------------
        # OBJECT ENTITY
        # --------------------------------------------------------------

        for name in names:

            value = getattr(
                entity,
                name,
                None,
            )

            # Boolean values such as:
            #
            #     evidence=False
            #
            # must not become:
            #
            #     "False"
            #
            if isinstance(value, bool):
                continue

            if (
                value is not None
                and str(value).strip()
            ):

                return str(
                    value
                ).strip()

        return ""
    # ==================================================================
    # STRUCTURED RESUME EVIDENCE
    # ==================================================================

    def _build_structured_resume_evidence(self, resume_path: Path) -> dict[str, Any]:
        """Reuse the existing parser-layer non-ontology extractors."""
        try:
            sections = self.resume_parser.parse(str(resume_path))
            resume = self.resume_builder.build(
                sections,
                source_file=str(resume_path),
                source_format="docx",
            )
            return {
                "education": self._json_safe(getattr(resume, "education", [])),
                "experience": self._json_safe(getattr(resume, "experience", [])),
                "languages": self._json_safe(getattr(resume, "languages", [])),
                "projects": self._json_safe(getattr(resume, "projects", [])),
                "awards": self._json_safe(getattr(resume, "awards", [])),
                "references": self._json_safe(getattr(resume, "references", [])),
                "contact": self._json_safe(getattr(getattr(resume, "personal_information", None), "__dict__", {})),
            }
        except Exception as exc:
            logger.warning(
                "Structured resume extraction unavailable; continuing with ontology evidence: %s",
                exc,
            )
            return {
                "education": [], "experience": [], "languages": [],
                "projects": [], "awards": [], "references": [], "contact": {},
                "error": f"{type(exc).__name__}: {exc}",
            }

    @classmethod
    def _experience_role_match(
        cls,
        *,
        subject: str,
        structured_resume: dict[str, Any],
        resume_text: str = "",
    ) -> dict[str, Any] | None:
        """Match role-family experience requirements against job titles.

        This is deterministic and explainable.  It intentionally uses the
        structured employment title as the primary evidence rather than
        claiming that a generic word overlap proves management experience.
        """
        normalized = cls._normalize(subject)
        if not normalized:
            return None

        role_family = None
        if re.search(r"\bsupervisory\b|\bsupervisor\b", normalized) and re.search(
            r"\bmanagerial\b|\bmanager\b|\bmanagement\b", normalized
        ):
            role_family = "managerial_supervisory"
        elif re.search(r"\bmanagerial\b|\bmanagement\b|\bmanager\b", normalized):
            role_family = "managerial"
        elif re.search(r"\bsupervisory\b|\bsupervisor\b", normalized):
            role_family = "supervisory"
        elif re.search(r"\bleadership\b|\blead\b|\bteam management\b", normalized):
            role_family = "leadership"
        elif re.search(r"\bexecutive\b|\bdirector\b|\bchief\b", normalized):
            role_family = "executive"
        else:
            return None

        records = structured_resume.get("experience", []) or []
        matches: list[tuple[float, Any, str, float]] = []

        # The parser's canonical field is ``title``, but enterprise resume
        # payloads can arrive through compatibility/legacy serializers with
        # aliases such as position, job_title, designation, or role. Accept
        # those aliases here so non-ontology experience evidence remains
        # useful even when a serializer changes shape.
        title_keys = (
            "title",
            "job_title",
            "position",
            "position_title",
            "designation",
            "role",
            "job_role",
        )

        for record in records:
            title = ""
            if isinstance(record, dict):
                for key in title_keys:
                    candidate_title = cls._flatten_structured_value(record.get(key, ""))
                    if candidate_title:
                        title = candidate_title
                        break
                if not title:
                    title = cls._flatten_structured_value(record.get("raw_header", ""))
            else:
                for key in title_keys:
                    candidate_title = cls._flatten_structured_value(getattr(record, key, ""))
                    if candidate_title:
                        title = candidate_title
                        break
                if not title:
                    title = cls._flatten_structured_value(getattr(record, "raw_header", ""))
            company = cls._flatten_structured_value(
                record.get("company", "") if isinstance(record, dict) else getattr(record, "company", "")
            )
            if not title:
                continue

            title_norm = cls._normalize(title)
            matched_phrase = None
            for phrase in cls.EXPERIENCE_ROLE_FAMILIES[role_family]:
                phrase_norm = cls._normalize(phrase)
                if phrase_norm and phrase_norm in title_norm:
                    matched_phrase = phrase
                    break

            if matched_phrase is None:
                continue

            duration = 0.0
            raw_duration = record.get("duration", 0.0) if isinstance(record, dict) else getattr(record, "duration", 0.0)
            try:
                duration = max(0.0, float(raw_duration or 0.0))
            except (TypeError, ValueError):
                duration = 0.0

            if duration <= 0.0:
                start = record.get("start_year", 0) if isinstance(record, dict) else getattr(record, "start_year", 0)
                end = record.get("end_year", 0) if isinstance(record, dict) else getattr(record, "end_year", 0)
                current = record.get("current_job", False) if isinstance(record, dict) else getattr(record, "current_job", False)
                try:
                    start_i = int(start or 0)
                    end_i = int(end or 0)
                    if start_i > 0:
                        if current:
                            from datetime import datetime
                            end_i = datetime.now().year
                        if end_i >= start_i:
                            duration = float(end_i - start_i)
                except (TypeError, ValueError):
                    duration = 0.0

            evidence = f"Structured resume evidence: {title}"
            if company:
                evidence += f" at {company}"
            evidence += f" — qualifies as a {role_family.replace('_', '/')} role."

            matches.append((1.0, record, evidence, duration))

        # Raw resume text is a secondary evidence source. It is intentionally
        # exact-title based: seeing "Store Manager" or "Managing Director" in
        # the resume is sufficient to establish role-family evidence, but we do
        # not infer management experience merely from generic verbs such as
        # "managed" or "supervised".
        if not matches and resume_text:
            text_norm = cls._normalize(resume_text)
            for phrase in cls.EXPERIENCE_ROLE_FAMILIES[role_family]:
                phrase_norm = cls._normalize(phrase)
                if not phrase_norm:
                    continue
                if re.search(r"\b" + re.escape(phrase_norm).replace(r"\ ", r"\s+") + r"\b", text_norm):
                    matches.append(
                        (
                            0.98,
                            None,
                            f"Resume text evidence: {phrase.title()} appears as a professional role, qualifying for the {role_family.replace('_', '/')} requirement.",
                            0.0,
                        )
                    )

        if not matches:
            return None

        qualifying_years = sum(item[3] for item in matches)
        evidence_lines = [item[2] for item in matches]
        return {
            "score": 1.0,
            "confidence": 0.98,
            "evidence": " ".join(evidence_lines),
            "qualifying_years": qualifying_years if qualifying_years > 0 else None,
            "role_family": role_family,
        }

    def _structured_resume_candidates(
        self,
        *,
        requirement_type: str,
        subject: str,
        structured_resume: dict[str, Any],
    ) -> list[dict[str, Any]]:
        """Truth-safe matching against structural resume evidence."""
        buckets = {
            "education": structured_resume.get("education", []),
            "experience": structured_resume.get("experience", []),
            "language": structured_resume.get("languages", []),
        }
        items = buckets.get(requirement_type, [])
        if not items:
            return []
        target = self._normalize(subject)
        if not target:
            return []
        target_tokens = set(target.split())
        results = []
        for item in items:
            text = self._flatten_structured_value(item)
            candidate = self._normalize(text)
            if not candidate:
                continue
            tokens = set(candidate.split())
            exact = target == candidate
            substring = target in candidate or candidate in target
            overlap = len(target_tokens & tokens) / max(len(target_tokens | tokens), 1)
            score = 1.0 if exact else 0.88 if substring else 0.62 if overlap >= 0.5 else 0.0
            if score <= 0:
                continue
            results.append({
                "score": score,
                "confidence": min(0.95, 0.75 + score * 0.20),
                "evidence": f"Structured resume evidence: {text[:500]}",
            })
        results.sort(key=lambda item: item["score"], reverse=True)
        return results

    @staticmethod
    def _flatten_structured_value(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, dict):
            allowed = {
                "raw", "description", "responsibilities", "achievements",
                "keywords", "skills", "technologies", "degree", "major",
                "institution", "level", "name", "proficiency", "title",
                "company", "industry", "location",
            }
            return " ".join(
                CandidateAnalysisService._flatten_structured_value(item)
                for key, item in value.items()
                if key in allowed
            )
        if isinstance(value, (list, tuple, set)):
            return " ".join(
                CandidateAnalysisService._flatten_structured_value(item)
                for item in value
            )
        return str(value)

    @staticmethod
    def _priority_weight(
        priority: str,
    ) -> float:

        return {
            "required": 3.0,
            "high": 3.0,
            "preferred": 2.0,
            "medium": 2.0,
            "contextual": 1.0,
            "low": 1.0,
        }.get(
            priority,
            1.0,
        )

    @staticmethod
    def _resume_years(
        text: str,
    ) -> float | None:

        matches = re.findall(
            r"\b(\d+(?:\.\d+)?)\s*\+?\s*years?\b",
            text,
            flags=re.IGNORECASE,
        )

        if not matches:
            return None

        return max(
            float(item)
            for item in matches
        )

    @staticmethod
    def _unique(
        values: Any,
    ) -> list[str]:

        result: list[str] = []

        seen: set[str] = set()

        for value in values:

            value = str(
                value or ""
            ).strip()

            if not value:
                continue

            key = value.casefold()

            if key in seen:
                continue

            seen.add(
                key
            )

            result.append(
                value
            )

        return result

