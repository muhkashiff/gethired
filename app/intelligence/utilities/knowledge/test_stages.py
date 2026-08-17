from __future__ import annotations

"""
Enterprise Interpretation Flow Test
Enterprise V13

Purpose
-------

Verify the NEW Enterprise V13 semantic architecture:

    resume_original.docx
            ↓
        resume text
            ↓
    EnterpriseResumePipeline
            ↓
    KnowledgeDocument
            ↓
    KnowledgeSentence[]
            ↓
    KnowledgeFact[]
            ↓
    KnowledgeInterpretation
            ↓
    SemanticResolver
            ↓
    SemanticResolution
            ↓
    SemanticEntity[]
    StatementRelation[]
    SemanticDependency[]
    SemanticCluster[]

This test intentionally DOES NOT use:

    SemanticResult
    SemanticMetadata
    SemanticStatistics
    old BusinessStatement contracts
    old data/resume.txt
    old .app/uploads paths

The test is designed to show exactly where entities
are flowing through the new architecture.
"""

from pathlib import Path
from typing import Any


# =====================================================================
# PROJECT / FILE IMPORTS
# =====================================================================

from app.intelligence.utilities.knowledge.enterprise_resume_pipeline import (
    EnterpriseResumePipeline,
)

from app.intelligence.utilities.knowledge.knowledge_models.knowledge_models import (
    KnowledgeDocument,
    KnowledgeFact,
    KnowledgeSentence,
)

from app.intelligence.utilities.knowledge.knowledge_extraction.interpretation_builder import (
    InterpretationBuilder,
)

from app.intelligence.utilities.knowledge.semantic_reasoning.semantic_models import (
    SemanticEntity,
    SemanticDependency,
    SemanticCluster,
    SemanticResolution,
    StatementRelation,
)


# =====================================================================
# DOCX IMPORT
# =====================================================================

try:

    from docx import Document

except ImportError as exc:

    raise ImportError(
        "python-docx is required to read "

        "uploads/project_2/resume_original.docx"
    ) from exc


# =====================================================================
# CONFIGURATION
# =====================================================================

TEST_NAME = "ENTERPRISE INTERPRETATION FLOW TEST"

RESUME_RELATIVE_PATH = Path(
    "Uploads",
    "project_2",
    "resume_original.docx",
)


# =====================================================================
# PATH RESOLUTION
# =====================================================================

def find_project_root() -> Path:
    """
    Locate the project root dynamically.

    We do NOT assume the current working directory.

    The test file is somewhere under:

        gethired/app/intelligence/utilities/knowledge/

    The resume is at:

        gethired/Uploads/project2/resume_original.docx
    """

    current_file = Path(__file__).resolve()

    for parent in current_file.parents:

        candidate = (
            parent
            / RESUME_RELATIVE_PATH
        )

        if candidate.exists():

            return parent

    raise FileNotFoundError(
        "Could not locate project root containing: "
        f"{RESUME_RELATIVE_PATH}"
    )


PROJECT_ROOT = find_project_root()

RESUME_FILE = (
    PROJECT_ROOT
    / RESUME_RELATIVE_PATH
)


# =====================================================================
# PRINT HELPERS
# =====================================================================

def line(char: str = "-") -> None:

    print(
        char * 80
    )


def section(title: str) -> None:

    print()

    line("=")

    print(title)

    line("=")


def subsection(title: str) -> None:

    print()

    line("-")

    print(title)

    line("-")


def ok(message: str) -> None:

    print(
        f"[PASS] {message}"
    )


def fail(message: str) -> None:

    print(
        f"[FAIL] {message}"
    )


def info(
    label: str,
    value: Any,
) -> None:

    print(
        f"{label:<32}: {value}"
    )


# =====================================================================
# RESUME READER
# =====================================================================

def read_resume_docx(
    path: Path,
) -> str:
    """
    Read the actual DOCX resume.

    Paragraphs are preserved as separate lines.
    Tables are also included because resumes frequently
    store experience / skills information inside tables.
    """

    if not path.exists():

        raise FileNotFoundError(
            f"Resume file does not exist: {path}"
        )

    document = Document(
        str(path)
    )

    parts: list[str] = []

    # -------------------------------------------------------------
    # Paragraphs
    # -------------------------------------------------------------

    for paragraph in document.paragraphs:

        text = (
            paragraph.text
            or ""
        ).strip()

        if text:

            parts.append(
                text
            )

    # -------------------------------------------------------------
    # Tables
    # -------------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            row_values = []

            for cell in row.cells:

                text = (
                    cell.text
                    or ""
                ).strip()

                if text:

                    row_values.append(
                        text
                    )

            if row_values:

                parts.append(
                    " | ".join(
                        row_values
                    )
                )

    return "\n".join(
        parts
    )


# =====================================================================
# SAFE ATTRIBUTE
# =====================================================================

def get_value(
    obj: Any,
    name: str,
    default: Any = None,
) -> Any:

    try:

        return getattr(
            obj,
            name,
            default,
        )

    except Exception:

        return default


# =====================================================================
# ENTITY CANONICAL NAME
# =====================================================================

def entity_name(
    entity: Any,
) -> str:

    canonical = get_value(
        entity,
        "canonical",
        "",
    )

    if canonical:

        return str(
            canonical
        )

    label = get_value(
        entity,
        "label",
        "",
    )

    if label:

        return str(
            label
        )

    normalized = get_value(
        entity,
        "normalized",
        "",
    )

    if normalized:

        return str(
            normalized
        )

    original = get_value(
        entity,
        "original",
        "",
    )

    return str(
        original
        or ""
    )


# =====================================================================
# ENTITY INSPECTION
# =====================================================================

def inspect_entities(
    entities: list[Any],
    title: str,
    limit: int = 30,
) -> None:

    subsection(
        title
    )

    info(
        "Entity count",
        len(
            entities
        ),
    )

    if not entities:

        print(
            "  NO ENTITIES"
        )

        return

    for index, entity in enumerate(
        entities[:limit],
        start=1,
    ):

        print()

        print(
            f"[{index}] "
            f"{entity.__class__.__name__}"
        )

        info(
            "entity_id",
            get_value(
                entity,
                "entity_id",
                "",
            ),
        )

        info(
            "entity_type",
            get_value(
                entity,
                "entity_type",
                "",
            ),
        )

        info(
            "canonical",
            entity_name(
                entity
            ),
        )

        info(
            "normalized",
            get_value(
                entity,
                "normalized",
                "",
            ),
        )

        info(
            "original",
            get_value(
                entity,
                "original",
                "",
            ),
        )

        info(
            "fact_id",
            get_value(
                entity,
                "fact_id",
                get_value(
                    entity,
                    "source_fact_id",
                    "",
                ),
            ),
        )

        info(
            "sentence_index",
            get_value(
                entity,
                "sentence_index",
                -1,
            ),
        )

        info(
            "confidence",
            get_value(
                entity,
                "confidence",
                0.0,
            ),
        )


# =====================================================================
# FACT INSPECTION
# =====================================================================

def inspect_facts(
    document: KnowledgeDocument,
    limit: int = 20,
) -> None:

    subsection(
        "KNOWLEDGE FACT FLOW"
    )

    facts = (
        document.facts
    )

    info(
        "Facts",
        len(
            facts
        ),
    )

    if not facts:

        fail(
            "KnowledgeDocument contains zero facts."
        )

        return

    for index, fact in enumerate(
        facts[:limit],
        start=1,
    ):

        print()

        print(
            f"[FACT {index}]"
        )

        info(
            "fact class",
            fact.__class__.__name__,
        )

        info(
            "fact id",
            get_value(
                fact,
                "fact_id",
                get_value(
                    fact,
                    "id",
                    "",
                ),
            ),
        )

        info(
            "fact text",
            get_value(
                fact,
                "text",
                "",
            ),
        )

        interpretation = get_value(
            fact,
            "interpretation",
            None,
        )

        if interpretation is None:

            fail(
                "Fact has NO interpretation."
            )

            continue

        ok(
            "Fact contains KnowledgeInterpretation."
        )

        entities = get_value(
            interpretation,
            "entities",
            [],
        ) or []

        info(
            "interpretation class",
            interpretation.__class__.__name__,
        )

        info(
            "interpretation entities",
            len(
                entities
            ),
        )

        for entity in entities[:10]:

            print(
                "    ENTITY:",
                entity_name(
                    entity
                ),
                "| TYPE:",
                get_value(
                    entity,
                    "entity_type",
                    "",
                ),
            )


# =====================================================================
# INTERPRETATION INSPECTION
# =====================================================================

def inspect_interpretations(
    document: KnowledgeDocument,
    limit: int = 20,
) -> int:
    """
    Inspect KnowledgeInterpretation objects already attached
    to KnowledgeFacts.

    Returns total number of extracted entities embedded
    inside interpretations.
    """

    subsection(
        "KNOWLEDGE INTERPRETATION FLOW"
    )

    total_interpretations = 0

    total_entities = 0

    for index, fact in enumerate(
        document.facts[:limit],
        start=1,
    ):

        interpretation = get_value(
            fact,
            "interpretation",
            None,
        )

        if interpretation is None:

            continue

        total_interpretations += 1

        entities = get_value(
            interpretation,
            "entities",
            [],
        ) or []

        total_entities += len(
            entities
        )

        print()

        print(
            f"[INTERPRETATION {index}]"
        )

        info(
            "fact",
            get_value(
                fact,
                "text",
                "",
            )[:150],
        )

        info(
            "interpretation class",
            interpretation.__class__.__name__,
        )

        info(
            "entity count",
            len(
                entities
            ),
        )

        if entities:

            for entity in entities[:10]:

                print(
                    "    ->",
                    entity.__class__.__name__,
                    "|",
                    get_value(
                        entity,
                        "entity_type",
                        "",
                    ),
                    "|",
                    entity_name(
                        entity
                    ),
                )

    info(
        "Interpretations checked",
        total_interpretations,
    )

    info(
        "Entities inside interpretations",
        total_entities,
    )

    return total_entities


# =====================================================================
# SEMANTIC RESOLUTION INSPECTION
# =====================================================================

def inspect_semantic_resolution(
    resolution: SemanticResolution,
) -> None:

    section(
        "SEMANTIC RESOLUTION"
    )

    info(
        "Resolution class",
        resolution.__class__.__name__,
    )

    info(
        "Semantic entities",
        resolution.entity_count,
    )

    info(
        "Statement relations",
        resolution.relation_count,
    )

    info(
        "Semantic dependencies",
        resolution.dependency_count,
    )

    info(
        "Semantic clusters",
        resolution.cluster_count,
    )

    info(
        "Business statements",
        resolution.statement_count,
    )

    info(
        "Fact count",
        resolution.fact_count,
    )

    info(
        "Sentence count",
        resolution.sentence_count,
    )

    info(
        "Confidence",
        resolution.confidence,
    )

    inspect_entities(
        resolution.entities,
        "SEMANTIC ENTITIES",
    )

    # -------------------------------------------------------------
    # Relations
    # -------------------------------------------------------------

    subsection(
        "STATEMENT RELATIONS"
    )

    info(
        "Count",
        resolution.relation_count,
    )

    for relation in resolution.relations[:20]:

        print(
            relation
        )

    # -------------------------------------------------------------
    # Dependencies
    # -------------------------------------------------------------

    subsection(
        "SEMANTIC DEPENDENCIES"
    )

    info(
        "Count",
        resolution.dependency_count,
    )

    for dependency in resolution.dependencies[:20]:

        print(
            dependency
        )

    # -------------------------------------------------------------
    # Clusters
    # -------------------------------------------------------------

    subsection(
        "SEMANTIC CLUSTERS"
    )

    info(
        "Count",
        resolution.cluster_count,
    )

    for cluster in resolution.clusters[:20]:

        print(
            cluster
        )


# =====================================================================
# RESOLUTION EXTRACTION FROM PIPELINE RESULT
# =====================================================================

def get_semantic_resolution(
    pipeline_result: Any,
) -> SemanticResolution | None:
    """
    Read the NEW semantic_resolution field.

    We intentionally do NOT use:

        semantic_result

    because that belongs to the old architecture.
    """

    resolution = get_value(
        pipeline_result,
        "semantic_resolution",
        None,
    )

    if isinstance(
        resolution,
        SemanticResolution,
    ):

        return resolution

    return None


# =====================================================================
# PIPELINE TEST
# =====================================================================

def test_interpretation_flow() -> None:

    section(
        TEST_NAME
    )

    print()

    print(
        "ARCHITECTURE:"
    )

    print(
        "DOCX"
    )

    print(
        "  ↓"
    )

    print(
        "KnowledgeDocument"
    )

    print(
        "  ↓"
    )

    print(
        "KnowledgeFact"
    )

    print(
        "  ↓"
    )

    print(
        "KnowledgeInterpretation"
    )

    print(
        "  ↓"
    )

    print(
        "SemanticResolver"
    )

    print(
        "  ↓"
    )

    print(
        "SemanticResolution"
    )

    print(
        "  ↓"
    )

    print(
        "SemanticEntity / Dependency / Cluster / Relation"
    )

    # =================================================================
    # 1. RESUME
    # =================================================================

    section(
        "1. RESUME INPUT"
    )

    info(
        "Project root",
        PROJECT_ROOT,
    )

    info(
        "Resume file",
        RESUME_FILE,
    )

    if not RESUME_FILE.exists():

        raise FileNotFoundError(
            f"Resume not found: {RESUME_FILE}"
        )

    resume_text = read_resume_docx(
        RESUME_FILE
    )

    if not resume_text.strip():

        raise ValueError(
            "resume_original.docx was read "
            "but produced empty text."
        )

    ok(
        "resume_original.docx loaded."
    )

    info(
        "Resume characters",
        len(
            resume_text
        ),
    )

    info(
        "Resume lines",
        len(
            resume_text.splitlines()
        ),
    )

    # =================================================================
    # 2. PIPELINE
    # =================================================================

    section(
        "2. ENTERPRISE PIPELINE"
    )

    pipeline = (
        EnterpriseResumePipeline()
    )

    pipeline_result = pipeline.run(
        resume_text
    )

    if pipeline_result is None:

        raise AssertionError(
            "EnterpriseResumePipeline.run() "
            "returned None."
        )

    info(
        "Pipeline result class",
        pipeline_result.__class__.__name__,
    )

    info(
        "Success",
        get_value(
            pipeline_result,
            "success",
            False,
        ),
    )

    info(
        "Failed stage",
        get_value(
            pipeline_result,
            "failed_stage",
            "",
        ),
    )

    error = get_value(
        pipeline_result,
        "error",
        None,
    )

    if error:

        print()

        fail(
            f"Pipeline error: {error!r}"
        )

    # =================================================================
    # 3. KNOWLEDGE DOCUMENT
    # =================================================================

    section(
        "3. KNOWLEDGE DOCUMENT"
    )

    document = get_value(
        pipeline_result,
        "knowledge_document",
        None,
    )

    if not isinstance(
        document,
        KnowledgeDocument,
    ):

        fail(
            "KnowledgeDocument was not created."
        )

        if error:

            print(
                "Pipeline error:",
                repr(
                    error
                ),
            )

        raise AssertionError(
            "KnowledgeDocument was not created."
        )

    ok(
        "KnowledgeDocument created."
    )

    info(
        "Document class",
        document.__class__.__name__,
    )

    info(
        "Sentences",
        len(
            document.sentences
        ),
    )

    info(
        "Facts",
        len(
            document.facts
        ),
    )

    # =================================================================
    # 4. FACT FLOW
    # =================================================================

    section(
        "4. KNOWLEDGE FACT FLOW"
    )

    inspect_facts(
        document
    )

    if not document.facts:

        raise AssertionError(
            "KnowledgeDocument contains zero facts."
        )

    # =================================================================
    # 5. INTERPRETATION FLOW
    # =================================================================

    section(
        "5. KNOWLEDGE INTERPRETATION FLOW"
    )

    interpretation_entities = (
        inspect_interpretations(
            document
        )
    )

    if interpretation_entities == 0:

        fail(
            "No entities exist inside "
            "KnowledgeInterpretation objects."
        )

    else:

        ok(
            f"{interpretation_entities} entities "
            "found inside interpretations."
        )

    # =================================================================
    # 6. SEMANTIC RESOLUTION
    # =================================================================

    section(
        "6. SEMANTIC RESOLUTION"
    )

    semantic_resolution = (
        get_semantic_resolution(
            pipeline_result
        )
    )

    if semantic_resolution is None:

        fail(
            "Enterprise pipeline did not expose "
            "a SemanticResolution object."
        )

        print()

        print(
            "Expected current architecture:"
        )

        print(
            "pipeline_result.semantic_resolution"
        )

        print()

        print(
            "Available result attributes:"
        )

        try:

            print(
                sorted(
                    vars(
                        pipeline_result
                    ).keys()
                )
            )

        except Exception:

            print(
                "Unable to inspect result attributes."
            )

        raise AssertionError(
            "SemanticResolution was not created."
        )

    ok(
        "SemanticResolution created."
    )

    # =================================================================
    # 7. SEMANTIC ENTITY FLOW
    # =================================================================

    inspect_semantic_resolution(
        semantic_resolution
    )

    if semantic_resolution.entity_count == 0:

        fail(
            "SemanticResolver produced zero "
            "SemanticEntity objects."
        )

        print()

        print(
            "THIS IS THE IMPORTANT FAILURE POINT."
        )

        print(
            "KnowledgeFact/KnowledgeInterpretation "
            "exists, but entities are not reaching "
            "SemanticResolution."
        )

        raise AssertionError(
            "Semantic resolution produced zero entities."
        )

    ok(
        f"SemanticResolver produced "
        f"{semantic_resolution.entity_count} entities."
    )

    # =================================================================
    # 8. ENTITY FLOW VALIDATION
    # =================================================================

    section(
        "7. ENTITY FLOW VALIDATION"
    )

    print()

    info(
        "KnowledgeDocument facts",
        len(
            document.facts
        ),
    )

    info(
        "Interpretation entities",
        interpretation_entities,
    )

    info(
        "Semantic entities",
        semantic_resolution.entity_count,
    )

    # -------------------------------------------------------------
    # Fact → Interpretation
    # -------------------------------------------------------------

    if interpretation_entities > 0:

        ok(
            "KnowledgeFact → KnowledgeInterpretation "
            "entity flow is working."
        )

    else:

        fail(
            "KnowledgeFact → KnowledgeInterpretation "
            "entity flow failed."
        )

    # -------------------------------------------------------------
    # Interpretation → Semantic
    # -------------------------------------------------------------

    if semantic_resolution.entity_count > 0:

        ok(
            "KnowledgeInterpretation → SemanticEntity "
            "flow is working."
        )

    else:

        fail(
            "KnowledgeInterpretation → SemanticEntity "
            "flow failed."
        )

    # =================================================================
    # 9. FINAL SUMMARY
    # =================================================================

    section(
        "FINAL INTEGRATION SUMMARY"
    )

    print()

    info(
        "Resume loaded",
        "PASS",
    )

    info(
        "KnowledgeDocument",
        "PASS",
    )

    info(
        "KnowledgeFacts",
        len(
            document.facts
        ),
    )

    info(
        "Interpretation entities",
        interpretation_entities,
    )

    info(
        "Semantic entities",
        semantic_resolution.entity_count,
    )

    info(
        "Statement relations",
        semantic_resolution.relation_count,
    )

    info(
        "Semantic dependencies",
        semantic_resolution.dependency_count,
    )

    info(
        "Semantic clusters",
        semantic_resolution.cluster_count,
    )

    print()

    print(
        "=" * 80
    )

    print(
        "ENTERPRISE INTERPRETATION FLOW TEST PASSED"
    )

    print(
        "=" * 80
    )
    

# =====================================================================
# MAIN
# =====================================================================

if __name__ == "__main__":

    test_interpretation_flow()