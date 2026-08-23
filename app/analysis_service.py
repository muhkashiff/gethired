"""
Candidate analysis orchestration.

This service is the application boundary between the Flask project workspace
and the existing intelligence pipeline.

Main flow:

    Resume DOCX + JD DOCX
            |
            v
    DocumentProcessingService (both documents)
            |
            v
    DocumentKnowledgeProfile + JDRequirementProfile
            |
            v
    Phase 3/4 matching objects when available
            |
            v
    Phase 5 ATS analyzer when available
            |
            v
    Candidate-facing result + repository-gap report

Important:
- No resume rewriting is performed here.
- Cover letters are intentionally outside this pipeline.
- The service never fabricates a missing skill as a candidate capability.
"""

from __future__ import annotations

import importlib
import json
import logging
import re
from dataclasses import asdict, is_dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from app.intelligence.utilities.knowledge.documents.document_input import (
    DocumentInput,
)
from app.intelligence.utilities.knowledge.documents.document_types import (
    DocumentType,
)
from app.intelligence.utilities.knowledge.documents.document_profile_builder import (
    DocumentProfileBuilder,
)
from app.intelligence.utilities.knowledge.documents.document_processing.document_processing_service import (
    DocumentProcessingService,
)
from app.intelligence.utilities.knowledge.jd_requirements.requirement_classifier import (
    JDRequirementClassifier,
)

logger = logging.getLogger("GetHired")


class AnalysisPipelineError(RuntimeError):
    """Raised when the candidate-analysis pipeline cannot complete."""


class CandidateAnalysisService:
    """Run and persist the candidate-vs-JD analysis."""

    def __init__(self, app_config: dict[str, Any]):
        self.upload_folder = Path(app_config["UPLOAD_FOLDER"])
        self.output_folder = Path(app_config["OUTPUT_FOLDER"])
        self.document_service = DocumentProcessingService()
        self.profile_builder = DocumentProfileBuilder()
        self.jd_classifier = JDRequirementClassifier()

    # ------------------------------------------------------------------
    # PUBLIC API
    # ------------------------------------------------------------------

    def run(self, project_id: int) -> dict[str, Any]:
        project_uploads = self.upload_folder / f"project_{project_id}"
        project_outputs = self.output_folder / f"project_{project_id}"
        project_outputs.mkdir(parents=True, exist_ok=True)

        resume_path = project_uploads / "resume_original.docx"
        jd_path = project_uploads / "job_description_original.docx"

        if not resume_path.exists():
            raise AnalysisPipelineError("Resume file is missing.")

        if not jd_path.exists():
            raise AnalysisPipelineError("Job Description file is missing.")

        logger.info("Analysis started | project=%s", project_id)

        resume_text = self._read_docx(resume_path)
        jd_text = self._read_docx(jd_path)

        resume_response = self.document_service.process(
            DocumentInput(
                text=resume_text,
                document_type=DocumentType.RESUME,
            )
        )

        if not resume_response.success:
            raise AnalysisPipelineError(
                f"Resume pipeline failed: {resume_response.error}"
            )

        jd_response = self.document_service.process(
            DocumentInput(
                text=jd_text,
                document_type=DocumentType.JD,
            )
        )

        if not jd_response.success:
            raise AnalysisPipelineError(
                f"Job Description pipeline failed: {jd_response.error}"
            )

        resume_profile = self.profile_builder.build(resume_response)
        jd_profile = self.profile_builder.build(jd_response)
        jd_requirements = self.jd_classifier.process(jd_profile)

        matching = self._build_matching(
            resume_profile=resume_profile,
            jd_profile=jd_profile,
            jd_requirements=jd_requirements,
            resume_text=resume_text,
        )

        ats = self._run_latest_ats_if_available(
            resume_text=resume_text,
            matching=matching,
        )

        result = self._build_candidate_result(
            project_id=project_id,
            resume_text=resume_text,
            jd_text=jd_text,
            resume_response=resume_response,
            jd_response=jd_response,
            resume_profile=resume_profile,
            jd_profile=jd_profile,
            jd_requirements=jd_requirements,
            matching=matching,
            ats=ats,
        )

        repository_gaps = self._collect_repository_gaps(
            resume_response.result,
            jd_response.result,
            matching.get("repository_gaps", []),
        )

        result["repository_gaps"] = repository_gaps
        result["completed_at"] = datetime.now(timezone.utc).isoformat()

        self._write_json(
            project_outputs / "analysis_result.json",
            result,
        )
        self._write_json(
            project_outputs / "repository_gaps.json",
            {
                "project_id": project_id,
                "generated_at": result["completed_at"],
                "gaps": repository_gaps,
            },
        )
        self._write_repository_gap_text(
            project_outputs / "repository_gaps.txt",
            project_id,
            repository_gaps,
        )

        logger.info(
            "Analysis completed | project=%s | overall=%s | ats=%s",
            project_id,
            result["scores"]["overall_match"],
            result["scores"]["ats_compatibility"],
        )

        return result

    def load(self, project_id: int) -> dict[str, Any] | None:
        path = (
            self.output_folder
            / f"project_{project_id}"
            / "analysis_result.json"
        )

        if not path.exists():
            return None

        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            logger.exception(
                "Could not load analysis result | project=%s",
                project_id,
            )
            return None

    # ------------------------------------------------------------------
    # DOCUMENT INPUT
    # ------------------------------------------------------------------

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            document = Document(str(path))
        except Exception as exc:
            raise AnalysisPipelineError(
                f"Could not read DOCX '{path.name}': {exc}"
            ) from exc

        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # Some resumes/JDs put important information in tables.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts).strip()

        if not text:
            raise AnalysisPipelineError(
                f"DOCX '{path.name}' contains no readable text."
            )

        return text

    # ------------------------------------------------------------------
    # MATCHING
    # ------------------------------------------------------------------

    def _build_matching(
        self,
        *,
        resume_profile: Any,
        jd_profile: Any,
        jd_requirements: Any,
        resume_text: str,
    ) -> dict[str, Any]:
        """Build Phase 3/4 objects when installed, with a safe local bridge."""

        resume_entities = list(
            getattr(getattr(resume_profile, "entities", None), "entities", [])
            or []
        )
        requirements = list(
            getattr(jd_requirements, "requirements", ()) or ()
        )

        records = []
        for requirement in requirements:
            records.append(
                self._match_requirement(
                    requirement,
                    resume_entities,
                    resume_text,
                )
            )

        # Try the typed Phase 3 -> Phase 4 objects from the user's newer
        # matching package. If that package is not present yet, the local
        # bridge remains the candidate-facing source of truth for this step.
        typed_profile = self._try_build_typed_match_profile(
            records=records,
            resume_profile=resume_profile,
            jd_requirements=jd_requirements,
        )

        strong = [item for item in records if item["status"] == "matched"]
        partial = [item for item in records if item["status"] == "partial"]
        missing = [item for item in records if item["status"] == "unmatched"]

        weighted_total = 0.0
        weight_total = 0.0
        for item in records:
            weight = self._priority_weight(item["priority"])
            weighted_total += item["score"] * weight
            weight_total += weight

        overall = (
            weighted_total / weight_total
            if weight_total
            else 0.0
        )

        evidence_values = [
            item["evidence_confidence"]
            for item in records
            if item["status"] != "unmatched"
        ]
        evidence_strength = (
            sum(evidence_values) / len(evidence_values)
            if evidence_values
            else 0.0
        )

        return {
            "requirements": records,
            "strong_matches": strong,
            "partial_matches": partial,
            "missing_matches": missing,
            "overall_match": round(overall * 100, 1),
            "evidence_strength": round(evidence_strength * 100, 1),
            "typed_profile": typed_profile,
            "repository_gaps": [],
        }

    def _try_build_typed_match_profile(
        self,
        *,
        records: list[dict[str, Any]],
        resume_profile: Any,
        jd_requirements: Any,
    ) -> Any:
        """
        Use the newer typed Phase 3/4 implementation when it exists.

        The current repository snapshot may not contain those modules yet,
        so ImportError is intentionally treated as a compatibility path.
        """
        try:
            match_models = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.match_models"
            )
            enrichment_models = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.enrichment_models"
            )
            enricher_module = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.match_enricher"
            )
            gap_module = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.gap_analyzer"
            )
            profile_builder_module = importlib.import_module(
                "app.intelligence.utilities.knowledge.matching.knowledge_match_profile_builder"
            )

            requirement_match_cls = match_models.RequirementMatch
            match_result_cls = match_models.KnowledgeMatchResult
            match_status = match_models.MatchStatus
            match_basis = match_models.MatchBasis

            typed_matches = []
            for record in records:
                status_name = record["status"]
                status = {
                    "matched": getattr(match_status, "MATCHED"),
                    "partial": getattr(match_status, "PARTIAL"),
                    "unmatched": getattr(match_status, "UNMATCHED"),
                }[status_name]

                basis = (
                    getattr(match_basis, "CANONICAL")
                    if status_name == "matched"
                    else (
                        getattr(match_basis, "CANONICAL")
                        if status_name == "partial"
                        else getattr(match_basis, "NONE")
                    )
                )

                typed_matches.append(
                    requirement_match_cls(
                        requirement_id=record["requirement_id"],
                        requirement_subject=record["subject"],
                        requirement_type=record["requirement_type"],
                        priority=record["priority"],
                        status=status,
                        score=record["score"],
                        basis=basis,
                        candidate_entity_ids=tuple(
                            record.get("candidate_entity_ids", [])
                        ),
                        candidate_evidence=tuple(
                            record.get("evidence", [])
                        ),
                        evidence_count=len(record.get("evidence", [])),
                    )
                )

            match_result = match_result_cls.from_matches(typed_matches)
            enriched = enricher_module.KnowledgeMatchEnricher().process(
                match_result=match_result,
                resume_profile=resume_profile,
                jd_requirement_profile=jd_requirements,
            )
            gap = gap_module.KnowledgeGapAnalyzer().process(enriched)
            profile = profile_builder_module.KnowledgeMatchProfileBuilder().process(
                match_result=match_result,
                enriched_match_result=enriched,
                gap_analysis_result=gap,
            )

            # Keep the exact Phase 4 object for Phase 5 if available.
            return profile

        except (ImportError, AttributeError, TypeError, ValueError) as exc:
            logger.info(
                "Typed Phase 3/4 matcher unavailable; using application bridge: %s",
                exc,
            )
            return None

    def _match_requirement(
        self,
        requirement: Any,
        resume_entities: list[Any],
        resume_text: str,
    ) -> dict[str, Any]:
        subject = str(getattr(requirement, "subject", "") or "").strip()
        requirement_type = self._enum_value(
            getattr(requirement, "requirement_type", "unknown")
        )
        priority = self._enum_value(
            getattr(requirement, "priority", "contextual")
        )

        subject_norm = self._normalize(subject)
        subject_tokens = set(subject_norm.split())

        candidates = []
        for entity in resume_entities:
            canonical = self._entity_text(
                entity,
                "canonical",
                "label",
                "normalized",
                "name",
                "original",
            )
            if not canonical:
                continue

            entity_norm = self._normalize(canonical)
            if not entity_norm:
                continue

            entity_tokens = set(entity_norm.split())
            exact = subject_norm == entity_norm
            substring = (
                subject_norm in entity_norm
                or entity_norm in subject_norm
            )
            overlap = (
                len(subject_tokens & entity_tokens)
                / max(len(subject_tokens | entity_tokens), 1)
            )

            score = 0.0
            basis = "none"
            if exact:
                score = 1.0
                basis = "canonical"
            elif substring:
                score = 0.82
                basis = "related"
            elif overlap >= 0.5:
                score = 0.65
                basis = "semantic_overlap"

            if score > 0:
                candidates.append((score, basis, entity))

        candidates.sort(key=lambda item: item[0], reverse=True)

        best = candidates[0] if candidates else None

        if best is None:
            status = "unmatched"
            score = 0.0
            basis = "none"
            candidate = None
        elif best[0] >= 0.8:
            status = "matched"
            score = best[0]
            basis = best[1]
            candidate = best[2]
        else:
            status = "partial"
            score = best[0]
            basis = best[1]
            candidate = best[2]

        evidence = []
        candidate_entity_ids = []
        evidence_confidence = 0.0

        if candidate is not None:
            candidate_id = self._entity_text(
                candidate,
                "entity_id",
                "id",
                "node_id",
            )
            if candidate_id:
                candidate_entity_ids.append(candidate_id)

            source_text = self._entity_text(
                candidate,
                "source_text",
                "description",
                "achievement",
                "evidence",
            )
            if source_text:
                evidence.append(source_text)

            try:
                evidence_confidence = float(
                    getattr(candidate, "confidence", 0.0)
                    if not isinstance(candidate, dict)
                    else candidate.get("confidence", 0.0)
                )
            except (TypeError, ValueError):
                evidence_confidence = 0.0

        # Experience requirements need a separate truth-safe observation.
        if requirement_type == "experience":
            minimum_years = getattr(requirement, "minimum_years", None)
            resume_years = self._resume_years(resume_text)
            if minimum_years is not None and resume_years is not None:
                if resume_years >= float(minimum_years):
                    status = "matched"
                    score = max(score, 1.0)
                    evidence.append(
                        f"Resume contains approximately {resume_years:g} years of stated experience."
                    )
                elif resume_years > 0:
                    status = "partial"
                    score = max(score, min(resume_years / float(minimum_years), 0.79))
                    evidence.append(
                        f"Resume states approximately {resume_years:g} years; JD asks for {float(minimum_years):g}+ years."
                    )

        return {
            "requirement_id": str(
                getattr(requirement, "requirement_id", subject)
            ),
            "subject": subject,
            "requirement_type": requirement_type,
            "priority": priority,
            "mandatory": bool(getattr(requirement, "mandatory", False)),
            "preferred": bool(getattr(requirement, "preferred", False)),
            "minimum_years": getattr(requirement, "minimum_years", None),
            "status": status,
            "score": round(score, 4),
            "basis": basis,
            "candidate_entity_ids": candidate_entity_ids,
            "evidence": evidence,
            "evidence_confidence": round(evidence_confidence, 4),
        }

    # ------------------------------------------------------------------
    # ATS
    # ------------------------------------------------------------------

    def _run_latest_ats_if_available(
        self,
        *,
        resume_text: str,
        matching: dict[str, Any],
    ) -> dict[str, Any]:
        profile = matching.get("typed_profile")
        if profile is not None:
            try:
                policy_module = importlib.import_module(
                    "app.intelligence.utilities.knowledge.ats.ats_analysis_policy"
                )
                analyzer_module = importlib.import_module(
                    "app.intelligence.utilities.knowledge.ats.ats_resume_analyzer"
                )
                request_module = importlib.import_module(
                    "app.intelligence.utilities.knowledge.ats.ats_analysis_request"
                )

                policy = policy_module.ATSAnalysisPolicy()
                analyzer = analyzer_module.ATSResumeAnalyzer(policy=policy)
                request = request_module.ATSResumeAnalysisRequest(
                    resume_text=resume_text,
                    knowledge_match_profile=profile,
                    policy=policy,
                )
                result = analyzer.process(request)

                score = float(
                    getattr(
                        getattr(result, "ats_score", None),
                        "score",
                        getattr(getattr(result, "score", None), "score", 0.0),
                    )
                )

                return {
                    "available": True,
                    "score": round(score * 100, 1),
                    "confidence": round(
                        float(getattr(result, "confidence", 0.0)) * 100,
                        1,
                    ),
                    "keyword_analysis": self._object_summary(
                        getattr(result, "keyword_analysis", None)
                    ),
                    "section_analysis": self._object_summary(
                        getattr(result, "section_analysis", None)
                    ),
                    "formatting_analysis": self._object_summary(
                        getattr(result, "formatting_analysis", None)
                    ),
                    "readability_analysis": self._object_summary(
                        getattr(result, "readability_analysis", None)
                    ),
                    "quantification_analysis": self._object_summary(
                        getattr(result, "quantification_analysis", None)
                    ),
                    "parseability_analysis": self._object_summary(
                        getattr(result, "parseability_analysis", None)
                    ),
                }
            except (ImportError, AttributeError, TypeError, ValueError) as exc:
                logger.info("Latest Phase 5 ATS unavailable: %s", exc)

        return self._fallback_ats(resume_text, matching)

    def _fallback_ats(
        self,
        resume_text: str,
        matching: dict[str, Any],
    ) -> dict[str, Any]:
        requirements = matching["requirements"]
        required_subjects = [
            item["subject"]
            for item in requirements
            if item["priority"] in {"required", "high"}
        ]
        if not required_subjects:
            required_subjects = [item["subject"] for item in requirements]

        lower = resume_text.casefold()
        matched_keywords = [
            subject
            for subject in required_subjects
            if subject.casefold() in lower
        ]
        missing_keywords = [
            subject
            for subject in required_subjects
            if subject not in matched_keywords
        ]

        keyword_score = (
            len(matched_keywords) / len(required_subjects)
            if required_subjects
            else 1.0
        )

        required_sections = (
            "summary",
            "experience",
            "skills",
            "education",
        )
        detected_sections = [
            section
            for section in required_sections
            if section in lower
        ]
        section_score = len(detected_sections) / len(required_sections)

        quantified = len(re.findall(r"\b\d+(?:\.\d+)?\s*%", resume_text))
        quantification_score = min(quantified / 3.0, 1.0)

        parseability_score = 1.0 if resume_text.strip() else 0.0

        score = (
            keyword_score * 0.55
            + section_score * 0.20
            + quantification_score * 0.15
            + parseability_score * 0.10
        )

        return {
            "available": False,
            "score": round(score * 100, 1),
            "confidence": 60.0,
            "keyword_analysis": {
                "required_keywords": required_subjects,
                "matched_keywords": matched_keywords,
                "missing_keywords": missing_keywords,
                "keyword_coverage_score": round(keyword_score * 100, 1),
            },
            "section_analysis": {
                "detected_sections": detected_sections,
                "missing_sections": [
                    item
                    for item in required_sections
                    if item not in detected_sections
                ],
                "section_completeness_score": round(section_score * 100, 1),
            },
            "quantification_analysis": {
                "quantified_achievement_count": quantified,
                "quantification_score": round(quantification_score * 100, 1),
            },
            "parseability_analysis": {
                "parseable": bool(resume_text.strip()),
                "parseability_score": round(parseability_score * 100, 1),
            },
        }

    # ------------------------------------------------------------------
    # RESULT
    # ------------------------------------------------------------------

    def _build_candidate_result(self, **kwargs: Any) -> dict[str, Any]:
        resume_profile = kwargs["resume_profile"]
        jd_profile = kwargs["jd_profile"]
        jd_requirements = kwargs["jd_requirements"]
        matching = kwargs["matching"]
        ats = kwargs["ats"]
        resume_response = kwargs["resume_response"]
        jd_response = kwargs["jd_response"]

        experience_gap = []
        for item in matching["requirements"]:
            if item["requirement_type"] == "experience":
                experience_gap.append(
                    {
                        "jd_requirement": item["subject"],
                        "minimum_years": item["minimum_years"],
                        "status": item["status"],
                        "evidence": item["evidence"],
                    }
                )

        recommendations = self._recommendations(
            matching=matching,
            ats=ats,
        )

        return {
            "project_id": kwargs["project_id"],
            "scores": {
                "overall_match": matching["overall_match"],
                "ats_compatibility": ats["score"],
                "evidence_strength": matching["evidence_strength"],
            },
            "matches": {
                "strong": matching["strong_matches"],
                "partial": matching["partial_matches"],
                "missing": matching["missing_matches"],
            },
            "experience_gap": experience_gap,
            "recommendations": recommendations,
            "categories": self._category_summary(resume_profile, jd_profile),
            "requirements_summary": {
                "total": len(getattr(jd_requirements, "requirements", ()) or ()),
                "required": getattr(jd_requirements, "required_count", 0),
                "preferred": getattr(jd_requirements, "preferred_count", 0),
                "contextual": getattr(jd_requirements, "contextual_count", 0),
            },
            "ats": ats,
            "pipeline": {
                "resume_entities": len(
                    getattr(
                        getattr(resume_profile, "entities", None),
                        "entities",
                        [],
                    )
                    or []
                ),
                "jd_entities": len(
                    getattr(
                        getattr(jd_profile, "entities", None),
                        "entities",
                        [],
                    )
                    or []
                ),
                "resume_semantic_entities": len(
                    getattr(resume_response, "semantic_entities", []) or []
                ),
                "jd_semantic_entities": len(
                    getattr(jd_response, "semantic_entities", []) or []
                ),
                "resume_business_statements": len(
                    getattr(resume_response, "business_statements", []) or []
                ),
                "jd_business_statements": len(
                    getattr(jd_response, "business_statements", []) or []
                ),
                "typed_phase4_profile": matching["typed_profile"] is not None,
            },
        }

    def _recommendations(
        self,
        *,
        matching: dict[str, Any],
        ats: dict[str, Any],
    ) -> list[str]:
        recommendations: list[str] = []

        for item in matching["strong_matches"][:3]:
            if item["evidence"]:
                recommendations.append(
                    f"Surface {item['subject']} evidence clearly in the resume."
                )

        for item in matching["partial_matches"][:4]:
            recommendations.append(
                f"Strengthen evidence for {item['subject']} if the experience is genuinely possessed."
            )

        for item in matching["missing_matches"][:5]:
            recommendations.append(
                f"Consider {item['subject']} only if genuinely possessed; do not add unsupported keywords."
            )

        missing_sections = (
            ats.get("section_analysis", {}) or {}
        ).get("missing_sections", [])
        if missing_sections:
            recommendations.append(
                "Improve resume structure by addressing missing ATS-relevant sections."
            )

        quantification = (
            ats.get("quantification_analysis", {}) or {}
        ).get("quantification_score", 100)
        if float(quantification or 0) < 70:
            recommendations.append(
                "Add measurable outcomes to genuine achievements where the resume currently lacks quantified evidence."
            )

        if not recommendations:
            recommendations.append(
                "The current analysis found no high-confidence improvement recommendation beyond maintaining evidence accuracy."
            )

        # De-duplicate while preserving order.
        return list(dict.fromkeys(recommendations))

    # ------------------------------------------------------------------
    # CATEGORIES / ADMIN GAPS
    # ------------------------------------------------------------------

    def _category_summary(self, resume_profile: Any, jd_profile: Any) -> dict[str, list[str]]:
        entities = list(
            getattr(getattr(resume_profile, "entities", None), "entities", [])
            or []
        )
        jd_entities = list(
            getattr(getattr(jd_profile, "entities", None), "entities", [])
            or []
        )

        groups = {
            "skills": {"skill"},
            "standards": {"standard"},
            "certifications": {"certification"},
            "methodologies": {"methodology"},
            "technologies": {"technology"},
            "actions": {"action"},
            "domains": {"domain"},
            "metrics": {"metric", "kpi"},
        }

        output: dict[str, list[str]] = {}
        for label, types in groups.items():
            output[label] = sorted(
                self._unique(
                    self._entity_text(entity, "canonical", "label", "name", "normalized")
                    for entity in entities
                    if self._enum_value(
                        self._entity_text(entity, "entity_type", "type", "category")
                    ) in types
                )
            )

        output["jd_requirements"] = sorted(
            self._unique(
                self._entity_text(entity, "canonical", "label", "name", "normalized")
                for entity in jd_entities
            )
        )
        return output

    def _collect_repository_gaps(self, *sources: Any) -> list[dict[str, Any]]:
        keys = {
            "unknown_terms",
            "unknown_entities",
            "unmatched_terms",
            "unrecognized_terms",
            "repository_gaps",
            "ontology_gaps",
        }
        found: list[dict[str, Any]] = []

        def walk(value: Any, path: str = "root") -> None:
            if value is None:
                return
            if isinstance(value, dict):
                for key, item in value.items():
                    key_text = str(key).casefold()
                    if key_text in keys:
                        self._append_gap_values(found, item, path + "." + str(key))
                    else:
                        walk(item, path + "." + str(key))
                return
            if isinstance(value, (list, tuple, set)):
                for index, item in enumerate(value):
                    walk(item, f"{path}[{index}]")
                return
            if is_dataclass(value):
                walk(asdict(value), path)
                return
            if hasattr(value, "__dict__"):
                try:
                    walk(vars(value), path)
                except TypeError:
                    pass

        for source in sources:
            walk(source)

        unique = {}
        for item in found:
            key = (
                item.get("term", ""),
                item.get("entity_type", ""),
            )
            unique[key] = item
        return sorted(unique.values(), key=lambda item: item["term"].casefold())

    @staticmethod
    def _append_gap_values(target: list[dict[str, Any]], value: Any, source: str) -> None:
        if isinstance(value, str):
            if value.strip():
                target.append({"term": value.strip(), "entity_type": "unknown", "source": source})
            return
        if isinstance(value, dict):
            term = (
                value.get("term")
                or value.get("label")
                or value.get("canonical")
                or value.get("text")
                or value.get("name")
            )
            if term:
                target.append({
                    "term": str(term),
                    "entity_type": str(value.get("entity_type") or value.get("type") or "unknown"),
                    "source": source,
                })
            else:
                for item in value.values():
                    CandidateAnalysisService._append_gap_values(target, item, source)
            return
        if isinstance(value, (list, tuple, set)):
            for item in value:
                CandidateAnalysisService._append_gap_values(target, item, source)

    # ------------------------------------------------------------------
    # SERIALIZATION
    # ------------------------------------------------------------------

    @staticmethod
    def _object_summary(value: Any) -> dict[str, Any]:
        if value is None:
            return {}
        if is_dataclass(value):
            return CandidateAnalysisService._json_safe(asdict(value))
        if isinstance(value, dict):
            return CandidateAnalysisService._json_safe(value)
        if hasattr(value, "__dict__"):
            return CandidateAnalysisService._json_safe(vars(value))
        return {"value": CandidateAnalysisService._json_safe(value)}

    @staticmethod
    def _json_safe(value: Any) -> Any:
        if value is None or isinstance(value, (str, int, float, bool)):
            return value
        if isinstance(value, dict):
            return {
                str(key): CandidateAnalysisService._json_safe(item)
                for key, item in value.items()
            }
        if isinstance(value, (list, tuple, set)):
            return [CandidateAnalysisService._json_safe(item) for item in value]
        if hasattr(value, "value") and not isinstance(value, (str, bytes)):
            try:
                return CandidateAnalysisService._json_safe(value.value)
            except Exception:
                pass
        if is_dataclass(value):
            return CandidateAnalysisService._json_safe(asdict(value))
        if hasattr(value, "__dict__"):
            return CandidateAnalysisService._json_safe(vars(value))
        return str(value)

    @staticmethod
    def _write_json(path: Path, value: Any) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            json.dumps(
                CandidateAnalysisService._json_safe(value),
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

    @staticmethod
    def _write_repository_gap_text(
        path: Path,
        project_id: int,
        gaps: list[dict[str, Any]],
    ) -> None:
        lines = [
            "GETHIRED REPOSITORY / ONTOLOGY GAP REPORT",
            "===========================================",
            f"Project: {project_id}",
            f"Generated: {datetime.now(timezone.utc).isoformat()}",
            "",
        ]

        if not gaps:
            lines.append("No explicit repository/ontology gaps were reported by the pipeline.")
        else:
            for gap in gaps:
                lines.append(
                    f"- {gap.get('term', '')} | type={gap.get('entity_type', 'unknown')} | source={gap.get('source', '')}"
                )

        path.write_text("\n".join(lines), encoding="utf-8")

    # ------------------------------------------------------------------
    # HELPERS
    # ------------------------------------------------------------------

    @staticmethod
    def _enum_value(value: Any) -> str:
        return str(getattr(value, "value", value)).strip().casefold()

    @staticmethod
    def _normalize(value: str) -> str:
        value = str(value or "").casefold()
        value = re.sub(r"[^a-z0-9+#.&/-]+", " ", value)
        return re.sub(r"\s+", " ", value).strip()

    @staticmethod
    def _entity_text(entity: Any, *names: str) -> str:
        if isinstance(entity, dict):
            for name in names:
                value = entity.get(name)
                if value is not None and str(value).strip():
                    return str(value).strip()
            return ""

        for name in names:
            value = getattr(entity, name, None)
            if value is not None and str(value).strip():
                return str(value).strip()
        return ""

    @staticmethod
    def _priority_weight(priority: str) -> float:
        return {
            "required": 3.0,
            "high": 3.0,
            "preferred": 2.0,
            "medium": 2.0,
            "contextual": 1.0,
            "low": 1.0,
        }.get(priority, 1.0)

    @staticmethod
    def _resume_years(text: str) -> float | None:
        matches = re.findall(
            r"\b(\d+(?:\.\d+)?)\s*\+?\s*years?\b",
            text,
            flags=re.IGNORECASE,
        )
        if not matches:
            return None
        return max(float(item) for item in matches)

    @staticmethod
    def _unique(values: Any) -> list[str]:
        result = []
        seen = set()
        for value in values:
            value = str(value or "").strip()
            if not value:
                continue
            key = value.casefold()
            if key in seen:
                continue
            seen.add(key)
            result.append(value)
        return result
