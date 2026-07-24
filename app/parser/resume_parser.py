"""
GetHired Resume Parser

Reads DOCX resumes and detects logical resume sections.
"""

from pathlib import Path
from docx import Document

from .section_detector import SectionDetector


class ResumeParser:
    """
    Production Resume Parser

    Stateless parser:
    parser = ResumeParser()
    sections = parser.parse("resume.docx")
    """

    def __init__(self):

        self.detector = SectionDetector()

    # ==========================================================
    # Read DOCX Paragraphs
    # ==========================================================

    def paragraphs(self, file_path):

        file_path = Path(file_path)

        document = Document(file_path)

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        return paragraphs

    # ==========================================================
    # Complete Resume Text
    # ==========================================================

    def full_text(self, file_path):

        return "\n".join(
            self.paragraphs(file_path)
        )

    # ==========================================================
    # Detect Resume Sections
    # ==========================================================

    def parse(self, file_path):

        paragraphs = self.paragraphs(file_path)

        return self.detector.detect(
            paragraphs
        )